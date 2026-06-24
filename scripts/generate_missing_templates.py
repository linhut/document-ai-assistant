# -*- coding: utf-8 -*-
"""生成缺失的 5 种公文类型 .dotx 模板文件"""

import os
import shutil
import zipfile
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# scripts/ 的上一级即项目根目录
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

# GB/T 9704 标准样式参数
STYLE_CONFIG = {
    "title": {
        "font_east_asia": "方正小标宋简体",
        "font_latin": "Times New Roman",
        "size": Pt(22),
        "bold": False,
        "alignment": WD_ALIGN_PARAGRAPH.CENTER,
        "space_after": Pt(20),
    },
    "body": {
        "font_east_asia": "仿宋_GB2312",
        "font_latin": "Times New Roman",
        "size": Pt(16),
        "alignment": WD_ALIGN_PARAGRAPH.JUSTIFY,
        "line_spacing": Pt(28.95),
        "first_line_indent": Cm(0.74),
    },
    "subtitle": {
        "font_east_asia": "黑体",
        "font_latin": "Times New Roman",
        "size": Pt(16),
        "bold": True,
        "alignment": WD_ALIGN_PARAGRAPH.LEFT,
        "line_spacing": Pt(28.95),
        "first_line_indent": Cm(0.74),
    },
    "signature": {
        "font_east_asia": "仿宋_GB2312",
        "font_latin": "Times New Roman",
        "size": Pt(16),
        "alignment": WD_ALIGN_PARAGRAPH.RIGHT,
    },
}

# 5 种需要补充的公文类型
TEMPLATES = {
    "minutes": {
        "cn_name": "纪要",
        "title": "XXX会议纪要",
        "paragraphs": [
            ("body", "时间：XXXX年XX月XX日"),
            ("body", "地点：XXX会议室"),
            ("body", "主持人：XXX"),
            ("body", "出席人员：XXX、XXX、XXX"),
            ("body", "缺席人员：XXX"),
            ("body", "记录人：XXX"),
            ("body", ""),
            ("body", "会议议定事项如下："),
            ("body", ""),
            ("subtitle", "一、关于XXX事项"),
            ("body", "会议认为，（具体内容）。"),
            ("body", "会议决定，（具体内容）。"),
            ("body", ""),
            ("subtitle", "二、关于XXX事项"),
            ("body", "会议认为，（具体内容）。"),
            ("body", "会议决定，（具体内容）。"),
            ("body", ""),
            ("body", ""),
            ("signature", "（单位名称）"),
            ("signature", "XXXX年XX月XX日"),
        ],
    },
    "instruction": {
        "cn_name": "指示",
        "title": "关于XXX工作的指示",
        "paragraphs": [
            ("body", "各部门、各单位："),
            ("body", ""),
            ("body", "当前，XXX工作面临新的形势和任务。为切实做好XXX工作，现作如下指示："),
            ("subtitle", "一、充分认识XXX工作的重要意义"),
            ("body", "（具体内容）"),
            ("subtitle", "二、明确XXX工作的总体要求和目标任务"),
            ("body", "（具体内容）"),
            ("subtitle", "三、切实加强XXX工作的组织领导"),
            ("body", "（具体内容）"),
            ("body", "各级各部门要认真贯彻落实本指示精神，确保各项工作落到实处。"),
            ("body", ""),
            ("body", ""),
            ("signature", "（单位名称）"),
            ("signature", "XXXX年XX月XX日"),
        ],
    },
    "regulation": {
        "cn_name": "制度",
        "title": "XXX管理办法",
        "paragraphs": [
            ("body", ""),
            ("subtitle", "第一章  总则"),
            ("body", ""),
            ("body", "第一条  为加强XXX管理，规范XXX行为，根据XXX有关规定，制定本办法。"),
            ("body", "第二条  本办法适用于XXX范围内的XXX活动。"),
            ("body", "第三条  XXX工作应当遵循XXX原则。"),
            ("body", ""),
            ("subtitle", "第二章  XXX"),
            ("body", ""),
            ("body", "第四条  （具体内容）"),
            ("body", "第五条  （具体内容）"),
            ("body", ""),
            ("subtitle", "第三章  XXX"),
            ("body", ""),
            ("body", "第六条  （具体内容）"),
            ("body", "第七条  （具体内容）"),
            ("body", ""),
            ("subtitle", "第四章  附则"),
            ("body", ""),
            ("body", "第八条  本办法由XXX负责解释。"),
            ("body", "第九条  本办法自XXXX年XX月XX日起施行。"),
            ("body", ""),
            ("body", ""),
            ("signature", "（单位名称）"),
            ("signature", "XXXX年XX月XX日"),
        ],
    },
    "summary": {
        "cn_name": "总结",
        "title": "关于XXX工作的总结",
        "paragraphs": [
            ("body", "XXX（上级机关）："),
            ("body", ""),
            ("body", "根据XXX要求，现将XXX工作情况总结如下："),
            ("subtitle", "一、基本情况"),
            ("body", "（总体概述工作背景和完成情况）"),
            ("subtitle", "二、主要做法和成效"),
            ("body", "（具体内容）"),
            ("subtitle", "三、存在的主要问题"),
            ("body", "（具体内容）"),
            ("subtitle", "四、下一步工作打算"),
            ("body", "（具体内容）"),
            ("body", ""),
            ("body", ""),
            ("signature", "（单位名称）"),
            ("signature", "XXXX年XX月XX日"),
        ],
    },
    "work_plan": {
        "cn_name": "工作方案",
        "title": "关于XXX工作的实施方案",
        "paragraphs": [
            ("body", ""),
            ("body", "为深入贯彻落实XXX精神，扎实推进XXX工作，制定本方案。"),
            ("body", ""),
            ("subtitle", "一、指导思想"),
            ("body", "（具体内容）"),
            ("body", ""),
            ("subtitle", "二、工作目标"),
            ("body", "（具体内容）"),
            ("body", ""),
            ("subtitle", "三、主要任务"),
            ("body", "（一）XXX。"),
            ("body", "（二）XXX。"),
            ("body", "（三）XXX。"),
            ("body", ""),
            ("subtitle", "四、实施步骤"),
            ("body", "（一）准备阶段（XXXX年XX月—XX月）。"),
            ("body", "（二）实施阶段（XXXX年XX月—XX月）。"),
            ("body", "（三）总结阶段（XXXX年XX月—XX月）。"),
            ("body", ""),
            ("subtitle", "五、保障措施"),
            ("body", "（具体内容）"),
            ("body", ""),
            ("body", ""),
            ("signature", "（单位名称）"),
            ("signature", "XXXX年XX月XX日"),
        ],
    },
}


def set_run_font(run, font_cfg):
    """设置 run 的四重字体属性"""
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_cfg["font_latin"])
    rFonts.set(qn("w:hAnsi"), font_cfg["font_latin"])
    rFonts.set(qn("w:eastAsia"), font_cfg["font_east_asia"])
    rFonts.set(qn("w:cs"), font_cfg["font_latin"])


def set_paragraph_spacing(paragraph, cfg):
    """设置段落行距和缩进"""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    if "line_spacing" in cfg:
        spacing.set(qn("w:line"), str(int(cfg["line_spacing"] * 20 / 12700 * 240)))
        spacing.set(qn("w:lineRule"), "exact")
    if "space_after" in cfg:
        spacing.set(qn("w:after"), str(int(cfg["space_after"] * 12700)))
    if "first_line_indent" in cfg:
        ind = pPr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            pPr.append(ind)
        ind.set(qn("w:firstLineChars"), "200")


def create_dotx(type_key, type_info, output_path):
    """创建 .dotx 模板文件"""
    doc = Document()

    # 页面设置 GB/T 9704
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    # 标题
    title_para = doc.add_paragraph()
    title_cfg = STYLE_CONFIG["title"]
    run = title_para.add_run(type_info["title"])
    set_run_font(run, title_cfg)
    run.font.size = title_cfg["size"]
    run.font.bold = title_cfg["bold"]
    title_para.alignment = title_cfg["alignment"]
    set_paragraph_spacing(title_para, title_cfg)

    # 正文段落
    for style_key, text in type_info["paragraphs"]:
        para = doc.add_paragraph()
        cfg = STYLE_CONFIG[style_key]
        if text:
            run = para.add_run(text)
            set_run_font(run, cfg)
            run.font.size = cfg["size"]
            if cfg.get("bold"):
                run.font.bold = True
        para.alignment = cfg.get("alignment", WD_ALIGN_PARAGRAPH.LEFT)
        set_paragraph_spacing(para, cfg)

    # 保存为 .docx 临时文件
    tmp_docx = output_path.replace(".dotx", "_tmp.docx")
    doc.save(tmp_docx)

    # 修改 Content Types 使其成为合法的 .dotx
    with zipfile.ZipFile(tmp_docx, "r") as zin:
        with zipfile.ZipFile(output_path, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "[Content_Types].xml":
                    data = data.replace(
                        b"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml",
                        b"application/vnd.ms-word.template.main+xml",
                    )
                zout.writestr(item, data)

    os.remove(tmp_docx)
    size = os.path.getsize(output_path)
    print(f"  OK {type_info['cn_name']}.dotx ({size:,} bytes)")


def main():
    print("=" * 60)
    print("generate missing .dotx template files")
    print("=" * 60)

    # 1. 生成到 公文模板/ 目录
    template_dir = os.path.join(PROJECT_ROOT, "公文模板")
    print(f"\nTarget: {template_dir}")
    for type_key, type_info in TEMPLATES.items():
        output = os.path.join(template_dir, f"{type_info['cn_name']}.dotx")
        create_dotx(type_key, type_info, output)

    # 2. 复制到 data/official_templates/（英文名）
    official_dir = os.path.join(PROJECT_ROOT, "data", "official_templates")
    print(f"\nTarget: {official_dir}")
    os.makedirs(official_dir, exist_ok=True)
    for type_key, type_info in TEMPLATES.items():
        src = os.path.join(template_dir, f"{type_info['cn_name']}.dotx")
        dst = os.path.join(official_dir, f"{type_key}.dotx")
        shutil.copy2(src, dst)
        size = os.path.getsize(dst)
        print(f"  OK {type_key}.dotx ({size:,} bytes)")

    print("\n" + "=" * 60)
    print("Done! 10 files generated (5 CN + 5 EN)")
    print("=" * 60)


if __name__ == "__main__":
    main()
