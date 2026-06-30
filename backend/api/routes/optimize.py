# This file is part of the Official Document AI Assistant.
# (c) 2026 Jose AI (https://www.linhut.cn)
# Licensed under the MIT License. See the LICENSE file for details.
"""
Optimize API routes: auto-fix and document generation.
"""
from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from starlette.background import BackgroundTask
from sqlalchemy.orm import Session
from pathlib import Path
import os
import asyncio

from db.database import get_db
from api.schemas.api_models import OptimizeRequest, OptimizeResponse
from services import document_service as svc
from services.document_service import _OPTIMIZED_SUFFIX
from utils.logger import logger

router = APIRouter()


# ---------------------------------------------------------------------------
#  Markdown 格式转换（前端实时预览用）— 必须在 /{doc_id} 之前定义
# ---------------------------------------------------------------------------

from pydantic import BaseModel
from typing import Any


class ParagraphData(BaseModel):
    text: str
    role: str | None = None
    is_heading: bool = False
    heading_level: int | None = None
    format: dict[str, Any] = {}


class MarkdownConvertRequest(BaseModel):
    paragraphs: list[ParagraphData]


# ---------------------------------------------------------------------------
#  Markdown 文本 AI 润色
# ---------------------------------------------------------------------------

class AIPolishRequest(BaseModel):
    text: str
    doc_type: str = "notice"


@router.post("/ai-polish")
async def ai_polish_markdown(body: AIPolishRequest, db: Session = Depends(get_db)):
    """
    AI 润色 Markdown 公文文本。

    调用已配置的 AI Provider 对用户输入的 Markdown 内容进行公文风格润色，
    返回润色后的文本。超时 15s，超时或无配置返回降级响应。
    """
    if not body.text.strip():
        raise HTTPException(status_code=400, detail="文本内容不能为空")

    # 获取 AI provider 配置（使用 FastAPI 注入的 db session）
    from db.models import AIConfig
    from ai.manager import create_provider, get_default_config
    from utils.crypto import decrypt_value

    provider_name = "openai"
    api_key = ""
    base_url = ""
    model = ""

    try:
        config = db.query(AIConfig).filter(AIConfig.is_active == True).first()
        if config:
            provider_name = config.provider
            api_key = decrypt_value(config.api_key_encrypted) or ""
            base_url = config.base_url or ""
            model = config.model or ""
        else:
            default = get_default_config()
            provider_name = default["provider"]
            api_key = default["api_key"]
            base_url = default["base_url"]
            model = default["model"]
    except Exception as e:
        logger.error(f"Failed to load AI config for polish: {e}")
        return {"success": False, "message": "AI 服务未配置，无法润色", "text": body.text}

    if not api_key:
        return {"success": False, "message": "AI API Key 未配置", "text": body.text}

    # 构建公文润色完整指令（作为 rewrite 的 text 参数）
    doc_type_label = {
        "notice": "通知", "request": "请示", "report": "报告",
        "letter": "函", "decision": "决定", "announcement": "通告",
        "notice_public": "公告", "meeting": "会议纪要",
    }.get(body.doc_type, "公文")

    full_prompt = f"""你是一位资深公文写作专家。请对以下 {doc_type_label} 内容进行润色，使其符合公文写作规范：
1. 语言正式、简洁、准确，符合公文语体
2. 逻辑清晰，层次分明
3. 保留原文的核心信息和结构（标题、正文、表格等）
4. 保留 Markdown 格式标记（#、**、| 等）
5. 不要添加原文没有的内容
6. 不要输出解释说明，只输出润色后的完整文本

原文：
{body.text}"""

    # 断路器：记录连续失败次数，超过阈值快速失败
    if not hasattr(ai_polish_markdown, '_fail_count'):
        ai_polish_markdown._fail_count = 0
        ai_polish_markdown._last_fail_ts = 0

    import time
    # 连续失败 ≥3 次，60 秒内快速失败
    if ai_polish_markdown._fail_count >= 3 and time.time() - ai_polish_markdown._last_fail_ts < 60:
        return {"success": False, "message": "AI 服务连续失败，请稍后再试", "text": body.text}

    ai_provider = None
    try:
        ai_provider = create_provider(provider_name, api_key, base_url, model)
        result = await asyncio.wait_for(
            ai_provider.rewrite(body.text, context=full_prompt),
            timeout=15.0,
        )
        # 成功：重置断路器
        ai_polish_markdown._fail_count = 0
        return {"success": True, "text": result, "provider": provider_name}
    except asyncio.TimeoutError:
        ai_polish_markdown._fail_count += 1
        ai_polish_markdown._last_fail_ts = time.time()
        logger.warning(f"AI polish timed out (15s), fail_count={ai_polish_markdown._fail_count}")
        return {"success": False, "message": "AI 润色超时，请稍后重试", "text": body.text}
    except Exception as e:
        ai_polish_markdown._fail_count += 1
        ai_polish_markdown._last_fail_ts = time.time()
        logger.error(f"AI polish failed (fail_count={ai_polish_markdown._fail_count}): {e}")
        return {"success": False, "message": f"AI 润色失败: {str(e)[:100]}", "text": body.text}
    finally:
        # 释放 httpx.AsyncClient 连接（熔断保护：即使 close() 失败也不阻塞响应）
        if ai_provider:
            try:
                await ai_provider.close()
            except Exception:
                pass


@router.post("/convert-markdown")
async def convert_markdown_text(body: MarkdownConvertRequest):
    """对段落文本执行 Markdown 格式识别与转换，返回转换后的段落。"""
    from core.document.models import (
        DocumentModel, DocumentMetadata, PageSetup,
        Paragraph, ParagraphFormat, Run, RunFormat,
    )
    from core.document.modifier import convert_markdown

    model = DocumentModel(
        metadata=DocumentMetadata(), page_setup=PageSetup(),
        paragraphs=[], tables=[], headers=[], footers=[],
    )

    for i, p in enumerate(body.paragraphs):
        rf = RunFormat(font_name=p.format.get('font_name'), font_size_pt=p.format.get('font_size_pt'), bold=p.format.get('bold'))
        pf = ParagraphFormat(alignment=p.format.get('alignment'), first_line_indent_pt=p.format.get('first_line_indent_pt'), line_spacing_pt=p.format.get('line_spacing_pt'))
        model.paragraphs.append(Paragraph(index=i, text=p.text, is_heading=p.is_heading, heading_level=p.heading_level, role=p.role, runs=[Run(index=0, text=p.text, format=rf)], format=pf))

    changes = convert_markdown(model)

    result = []
    for p in model.paragraphs:
        rf = p.runs[0].format if p.runs else RunFormat()
        result.append({
            "text": p.text, "role": p.role, "is_heading": p.is_heading, "heading_level": p.heading_level,
            "format": {"alignment": p.format.alignment, "first_line_indent_pt": p.format.first_line_indent_pt, "font_name": rf.font_name, "font_size_pt": rf.font_size_pt, "line_spacing_pt": p.format.line_spacing_pt, "bold": rf.bold},
        })

    # 序列化表格（markdown 表格转换后生成的 Table 对象）
    tables = []
    for t in model.tables:
        cells = []
        for c in t.cells:
            cell_paras = []
            for cp in c.paragraphs:
                rf = cp.runs[0].format if cp.runs else RunFormat()
                cell_paras.append({
                    "text": cp.text,
                    "format": {
                        "alignment": cp.format.alignment,
                        "font_name": rf.font_name,
                        "font_size_pt": rf.font_size_pt,
                        "bold": rf.bold,
                    },
                })
            cells.append({"row": c.row, "col": c.col, "text": c.text, "paragraphs": cell_paras})
        tables.append({"index": t.index, "rows": t.rows, "cols": t.cols, "cells": cells, "insert_after_index": t.insert_after_index})

    return {"success": True, "changes": changes, "paragraphs": result, "tables": tables}


# ---------------------------------------------------------------------------
#  Markdown 文本 → 预览数据（粘贴 Markdown + 选文种 → 生成公文预览）
# ---------------------------------------------------------------------------

import re as _re
import datetime as _dt


def _parse_margin(val):
    """解析 "3.7cm" → 37.0 (mm)"""
    if isinstance(val, (int, float)):
        return float(val) * 10
    s = str(val).strip().lower()
    m = _re.match(r'([\d.]+)\s*(cm|mm)?', s)
    if m:
        v = float(m.group(1))
        unit = m.group(2) or 'cm'
        return v * 10 if unit == 'cm' else v
    return 37.0


def _embed_stamp(
    output_path: str,
    image_base64: str,
    page_number: int = 0,
    x_mm: float = 30,
    y_mm: float = 25,
    width_mm: float = 40,
    height_mm: float = 40,
):
    """
    将印章图片嵌入到 docx 的指定页面。

    使用 python-docx 的 add_picture + 绝对定位实现。
    page_number=0 表示最后一页。
    """
    import base64
    import tempfile
    from docx import Document
    from docx.shared import Mm, Cm
    from docx.oxml.ns import qn
    from lxml import etree

    # 解码 base64 → 临时 PNG 文件
    # 支持 data:image/png;base64,... 格式
    b64_data = image_base64
    if ',' in b64_data:
        b64_data = b64_data.split(',', 1)[1]

    stamp_bytes = base64.b64decode(b64_data)
    stamp_tmp = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
    stamp_tmp.write(stamp_bytes)
    stamp_tmp.close()

    try:
        doc = Document(output_path)

        # 确定目标页面（page_number 为 0 或超出范围 → 最后一页）
        # python-docx 没有直接的分页 API，通过段落分页符定位
        paragraphs = doc.paragraphs
        if not paragraphs:
            return

        # 收集分页位置：在正文段落中查找分页符或估算页面
        page_breaks: list[int] = [0]  # 每个元素是页面起始段落索引
        for i, para in enumerate(paragraphs):
            # 检查段落是否有分页符
            pPr = para._element.find(qn('w:pPr'))
            if pPr is not None:
                # 段前分页 <w:pageBreakBefore/>
                page_break_before = pPr.find(qn('w:pageBreakBefore'))
                if page_break_before is not None:
                    page_breaks.append(i)
                    continue
                # 段内分页符 <w:r><w:br w:type="page"/></w:r>
                for run in para._element.findall(qn('w:r')):
                    br = run.find(qn('w:br'))
                    if br is not None and br.get(qn('w:type')) == 'page':
                        page_breaks.append(i + 1)
                        break

        # 目标页面索引（0-based）
        total_pages = len(page_breaks)
        target_page = page_number - 1 if page_number > 0 else total_pages - 1
        target_page = max(0, min(target_page, total_pages - 1))

        # 目标页面的第一个段落
        target_para_idx = page_breaks[target_page]
        target_para = paragraphs[target_para_idx]

        # 在目标段落前插入印章图片（使用绝对定位）
        # 通过 run 添加图片，然后用 OXML 设置绝对定位
        run = target_para.add_run()
        inline_shape = run.add_picture(
            stamp_tmp.name,
            width=Mm(width_mm),
            height=Mm(height_mm),
        )

        # 设置图片为绝对定位（Word 中印章浮于文字上方）
        # 通过修改 OXML 将 inline 转为 anchored
        _set_picture_absolute_position(
            inline_shape._inline,
            x_mm=x_mm,
            y_mm=y_mm,
            width_mm=width_mm,
            height_mm=height_mm,
        )

        doc.save(output_path)
    finally:
        os.unlink(stamp_tmp.name)


def _set_picture_absolute_position(
    inline_element,
    x_mm: float,
    y_mm: float,
    width_mm: float,
    height_mm: float,
):
    """
    将 python-docx 的 inline 图片转为绝对定位的 anchored 图片。

    使用 Word 的 wp:anchor 元素替代 wp:inline，
    实现印章浮于文字上方、固定于页面右下角的效果。
    """
    from docx.oxml.ns import qn
    from lxml import etree

    NSMAP = {
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
        'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
        'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
        'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    }

    drawing = inline_element.getparent()
    if drawing is None:
        return

    # 提取 inline 中的 extent 和 graphic 数据
    extent = inline_element.find(qn('wp:extent'))
    effectExtent = inline_element.find(qn('wp:effectExtent'))
    docPr = inline_element.find(qn('wp:docPr'))
    cNvGraphicFramePr = inline_element.find(qn('wp:cNvGraphicFramePr'))
    graphic = inline_element.find(qn('a:graphic'))

    cx = extent.get('cx') if extent is not None else str(int(width_mm * 36000))
    cy = extent.get('cy') if extent is not None else str(int(height_mm * 36000))

    # 创建 anchored 元素
    anchored = etree.SubElement(drawing, qn('wp:anchor'), nsmap=NSMAP)

    # 设置属性
    anchored.set('behindDoc', '0')
    anchored.set('distT', '0')
    anchored.set('distB', '0')
    anchored.set('distL', '0')
    anchored.set('distR', '0')
    anchored.set('simplePos', '0')
    anchored.set('relativeHeight', '251658240')  # 置于顶层
    anchored.set('locked', '0')
    anchored.set('layoutInCell', '1')
    anchored.set('allowOverlap', '1')

    # simplePos
    simplePos = etree.SubElement(anchored, qn('wp:simplePos'))
    simplePos.set('x', '0')
    simplePos.set('y', '0')

    # positionH — 相对于页面右侧
    posH = etree.SubElement(anchored, qn('wp:positionH'))
    posH.set('relativeFrom', 'page')
    posH_name = etree.SubElement(posH, qn('wp:posOffset'))
    posH_name.text = str(int(x_mm * 36000))

    # positionV — 相对于页面底部
    posV = etree.SubElement(anchored, qn('wp:positionV'))
    posV.set('relativeFrom', 'page')
    posV_name = etree.SubElement(posV, qn('wp:posOffset'))
    posV_name.text = str(int(y_mm * 36000))

    # extent
    ext = etree.SubElement(anchored, qn('wp:extent'))
    ext.set('cx', cx)
    ext.set('cy', cy)

    # effectExtent
    if effectExtent is not None:
        anchored.append(effectExtent)

    # wrapNone — 浮于文字上方
    wrapNone = etree.SubElement(anchored, qn('wp:wrapNone'))

    # docPr
    if docPr is not None:
        anchored.append(docPr)

    # cNvGraphicFramePr
    if cNvGraphicFramePr is not None:
        anchored.append(cNvGraphicFramePr)

    # graphic
    if graphic is not None:
        anchored.append(graphic)

    # 移除原 inline 元素
    drawing.remove(inline_element)


# ---------------------------------------------------------------------------
#  版头/版记直接注入（使用 python-docx API）
# ---------------------------------------------------------------------------

def _insert_before(new_para, reference_p, body) -> None:
    """将段落插入到 reference_p 之前。"""
    if reference_p is not None:
        body.insert(list(body).index(reference_p), new_para._element)
    else:
        body.append(new_para._element)


def _inject_header_to_docx(output_path: str, header_config: dict) -> None:
    """注入版头：发文机关标志 + 空行 + 发文字号 + 红色分隔线。

    使用 python-docx API 而非裸 OOXML，确保样式正确写入。
    公文版头顺序（从上到下）：机关名 → 空行 → 发文字号/签发人 → 红色分隔线
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from core.document.font_utils import set_run_font

        doc = Document(output_path)
        org_name = header_config.get('org_name', '')
        doc_number = header_config.get('doc_number', '')
        signer = header_config.get('signer', '')

        if not org_name:
            return

        body = doc.element.body
        first_p = body.find(qn('w:p'))

        # === 按正确顺序插入（从文档最前面开始） ===

        # 1. 发文机关标志：红色 30pt 方正小标宋 居中（最顶部）
        p_org = doc.add_paragraph()
        p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_org = p_org.add_run(org_name)
        set_run_font(run_org, '方正小标宋_GBK')
        run_org.font.size = Pt(30)
        run_org.font.color.rgb = RGBColor(0xE0, 0x00, 0x00)
        p_org.paragraph_format.line_spacing = Pt(28.95)
        p_org.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        _insert_before(p_org, first_p, body)

        # 2. 空二行
        for _ in range(2):
            p_empty = doc.add_paragraph()
            p_empty.paragraph_format.line_spacing = Pt(28.95)
            p_empty.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            _insert_before(p_empty, first_p, body)

        # 3. 发文字号（居中或左对齐）
        if doc_number:
            p_num = doc.add_paragraph()
            p_num.alignment = WD_ALIGN_PARAGRAPH.LEFT if signer else WD_ALIGN_PARAGRAPH.CENTER
            run_num = p_num.add_run(doc_number)
            set_run_font(run_num, '仿宋_GB2312')
            run_num.font.size = Pt(16)
            p_num.paragraph_format.line_spacing = Pt(28.95)
            p_num.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            _insert_before(p_num, first_p, body)

        # 4. 签发人行（右对齐）
        if signer:
            p_signer = doc.add_paragraph()
            p_signer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run_signer = p_signer.add_run(f'签发人：{signer}')
            set_run_font(run_signer, '仿宋_GB2312')
            run_signer.font.size = Pt(16)
            p_signer.paragraph_format.line_spacing = Pt(28.95)
            p_signer.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            _insert_before(p_signer, first_p, body)

        # 5. 红色分隔线（底部边框）— 紧贴在发文字号下方
        p_border = doc.add_paragraph()
        p_border.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pPr = p_border._element.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '15')
        bottom.set(qn('w:color'), 'E00000')
        bottom.set(qn('w:space'), '1')
        pBdr.append(bottom)
        pPr.append(pBdr)
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:line'), '60')  # 3pt 高度
        spacing.set(qn('w:lineRule'), 'exact')
        pPr.append(spacing)
        _insert_before(p_border, first_p, body)

        doc.save(output_path)
        logger.info(f"Injected header into {output_path}")
    except Exception as e:
        logger.error(f"Failed to inject header: {e}", exc_info=True)


def _inject_footer_to_docx(output_path: str, footer_config: dict) -> None:
    """注入版记：分隔线 + 抄送 + 印发 + 分隔线。

    使用 python-docx API 而非裸 OOXML，确保样式正确写入。
    GB/T 9704 §7.4：版记用四号仿宋（14pt），左空一字。
    """
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from core.document.font_utils import set_run_font

        doc = Document(output_path)
        cc = footer_config.get('cc', '')
        printer = footer_config.get('printer', '')
        print_date = footer_config.get('printDate', '')

        if not (cc or printer or print_date):
            return

        # 辅助：创建带顶部分隔线的段落
        def _add_border_para(doc, border_size='12'):
            p = doc.add_paragraph()
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            top = OxmlElement('w:top')
            top.set(qn('w:val'), 'single')
            top.set(qn('w:sz'), border_size)
            top.set(qn('w:color'), '000000')
            top.set(qn('w:space'), '1')
            pBdr.append(top)
            pPr.append(pBdr)
            # 设置段落高度
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:line'), '60')  # 3pt
            spacing.set(qn('w:lineRule'), 'exact')
            pPr.append(spacing)
            return p

        # 1. 粗分隔线（顶边框）
        _add_border_para(doc, '12')

        # 2. 抄送行（四号仿宋，左空一字）
        if cc:
            p_cc = doc.add_paragraph()
            p_cc.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_cc = p_cc.add_run(f'抄送：{cc}。')
            set_run_font(run_cc, '仿宋_GB2312')
            run_cc.font.size = Pt(14)
            p_cc.paragraph_format.line_spacing = Pt(28.95)
            p_cc.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            # 左空一字（约 14pt）
            p_cc.paragraph_format.left_indent = Pt(14)

            # 细分隔线（抄送与印发之间）
            if printer or print_date:
                _add_border_para(doc, '4')

        # 3. 印发行（四号仿宋，左空一字）
        if printer or print_date:
            p_info = doc.add_paragraph()
            p_info.alignment = WD_ALIGN_PARAGRAPH.LEFT
            parts = []
            if printer:
                parts.append(printer)
            if print_date:
                parts.append(f'{print_date}印发')
            run_info = p_info.add_run('        '.join(parts))
            set_run_font(run_info, '仿宋_GB2312')
            run_info.font.size = Pt(14)
            p_info.paragraph_format.line_spacing = Pt(28.95)
            p_info.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p_info.paragraph_format.left_indent = Pt(14)

        # 4. 底部粗分隔线
        _add_border_para(doc, '12')

        doc.save(output_path)
        logger.info(f"Injected footer into {output_path}")
    except Exception as e:
        logger.error(f"Failed to inject footer: {e}", exc_info=True)


def _make_simple_para(text: str, font_name: str, font_size_pt: int, align: str, color: str | None = None) -> 'OxmlElement':
    """创建一个简单的 OOXML 段落元素。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')

    # 对齐
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), align if align != 'justify' else 'both')
    pPr.append(jc)

    # 行距
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:line'), '579')  # 28.95pt
    spacing.set(qn('w:lineRule'), 'exact')
    pPr.append(spacing)

    p.append(pPr)

    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size_pt * 2)))  # half-points
    rPr.append(sz)
    if color:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color)
        rPr.append(c)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    p.append(r)

    return p


class MarkdownToPreviewRequest(BaseModel):
    markdown_text: str
    doc_type: str = "notice"


@router.post("/markdown-to-preview")
async def markdown_to_preview(body: MarkdownToPreviewRequest):
    """接收 Markdown 原始文本 + 文种，返回带版头版记的完整预览数据。"""
    from core.document.models import (
        DocumentModel, DocumentMetadata, PageSetup,
        Paragraph, ParagraphFormat, Run, RunFormat,
    )
    from core.document.modifier import convert_markdown
    from core.template.style_manager import get_template
    from core.rules.loader import load_rules_for_type

    # 1. 读取 YAML 规则
    merged_rules = load_rules_for_type(body.doc_type)

    # 2. 解析页边距
    margins = merged_rules.get('page_setup', {}).get('margins', {})
    page_setup = {
        "margin_top_mm": _parse_margin(margins.get('top', '3.7cm')),
        "margin_bottom_mm": _parse_margin(margins.get('bottom', '3.5cm')),
        "margin_left_mm": _parse_margin(margins.get('left', '2.8cm')),
        "margin_right_mm": _parse_margin(margins.get('right', '2.6cm')),
    }

    # 3. 构建版头段落
    header_paras = []
    header_fields = merged_rules.get('header_fields', {})

    org_name = header_fields.get('issuing_org', {}).get('default', '')
    if org_name:
        header_paras.append(Paragraph(
            index=len(header_paras), text=org_name,
            role='header_org', is_heading=False,
            runs=[Run(index=0, text=org_name, format=RunFormat(
                font_name='方正小标宋简体', font_size_pt=30.0, color='#E00000',
            ))],
            format=ParagraphFormat(alignment='center'),
        ))

    doc_number = header_fields.get('document_number', {}).get('default', '')
    if doc_number:
        header_paras.append(Paragraph(
            index=len(header_paras), text=doc_number,
            role='header_number', is_heading=False,
            runs=[Run(index=0, text=doc_number, format=RunFormat(
                font_name='仿宋_GB2312', font_size_pt=16.0,
            ))],
            format=ParagraphFormat(alignment='center'),
        ))

    # 红色分隔线标记（空文本，不写入 docx 内容）
    header_paras.append(Paragraph(
        index=len(header_paras), text='',
        role='red_line', is_heading=False,
        runs=[Run(index=0, text='', format=RunFormat())],
        format=ParagraphFormat(alignment='center'),
    ))

    # 4. 将 Markdown 文本按行拆分为段落
    lines = body.markdown_text.split('\n')

    # 5. 构建 DocumentModel（版头 + 正文）
    model = DocumentModel(
        metadata=DocumentMetadata(), page_setup=PageSetup(),
        paragraphs=list(header_paras),
        tables=[], headers=[], footers=[],
    )
    for i, line in enumerate(lines):
        idx = len(header_paras) + i
        rf = RunFormat(font_name='仿宋_GB2312', font_size_pt=16.0)
        pf = ParagraphFormat(alignment='justify', first_line_indent_pt=32.0, line_spacing_pt=28.95)
        model.paragraphs.append(Paragraph(
            index=idx, text=line, role='body', is_heading=False,
            runs=[Run(index=0, text=line, format=rf)], format=pf,
        ))

    # 6. 执行 Markdown 转换
    changes = convert_markdown(model)

    # 6.1 标题后插入空行（GB/T 9704：公文标题与正文之间应空一行）
    for i, para in enumerate(model.paragraphs):
        if para.is_heading and para.heading_level == 0 and para.role == 'title':
            # 在标题后插入一个空段落
            empty_para = Paragraph(
                index=i + 1, text='', role='body', is_heading=False,
                runs=[Run(index=0, text='', format=RunFormat())],
                format=ParagraphFormat(alignment='justify', line_spacing_pt=28.95),
            )
            model.paragraphs.insert(i + 1, empty_para)
            # 重建后续索引
            for j in range(i + 2, len(model.paragraphs)):
                model.paragraphs[j].index = j
            break

    # 7. 构建版记段落
    footer_paras = []
    footer_paras.append(Paragraph(
        index=len(model.paragraphs) + len(footer_paras),
        text='', role='footer_line', is_heading=False,
        runs=[Run(index=0, text='', format=RunFormat())],
        format=ParagraphFormat(alignment='center'),
    ))
    cc_text = header_fields.get('cc', {}).get('default', '')
    if cc_text:
        footer_paras.append(Paragraph(
            index=len(model.paragraphs) + len(footer_paras),
            text=f'抄送：{cc_text}', role='cc', is_heading=False,
            runs=[Run(index=0, text=f'抄送：{cc_text}', format=RunFormat(font_name='仿宋_GB2312', font_size_pt=14.0))],
            format=ParagraphFormat(alignment='left', left_indent_pt=16.0),
        ))
    # 印发机关 + 印发日期
    footer_fields = merged_rules.get('footer_fields', {})
    printer = footer_fields.get('printer', {}).get('default', '') or header_fields.get('issuing_org', {}).get('default', '')
    print_date = footer_fields.get('print_date', {}).get('default', '') or _dt.date.today().strftime('%Y年%m月%d日')
    if printer or print_date:
        footer_info_text = f'{printer}{"　" if printer and print_date else ""}{print_date}印发'
        footer_paras.append(Paragraph(
            index=len(model.paragraphs) + len(footer_paras),
            text=footer_info_text, role='footer_info', is_heading=False,
            runs=[Run(index=0, text=footer_info_text, format=RunFormat(font_name='仿宋_GB2312', font_size_pt=14.0))],
            format=ParagraphFormat(alignment='right'),
        ))
    model.paragraphs.extend(footer_paras)

    # 8. 获取模板信息
    template = get_template(body.doc_type)
    template_name = template.get("name", body.doc_type) if template else body.doc_type

    # 9. 序列化返回
    result_paras = []
    for p in model.paragraphs:
        rf = p.runs[0].format if p.runs else RunFormat()
        runs_data = [{"text": r.text or "", "bold": getattr(r.format, 'bold', None), "font_name": getattr(r.format, 'font_name', None), "font_size_pt": getattr(r.format, 'font_size_pt', None), "color": getattr(r.format, 'color', None)} for r in p.runs]
        result_paras.append({
            "text": p.text, "role": p.role, "is_heading": p.is_heading, "heading_level": p.heading_level,
            "format": {"alignment": p.format.alignment, "first_line_indent_pt": p.format.first_line_indent_pt, "left_indent_pt": getattr(p.format, 'left_indent_pt', None), "font_name": rf.font_name, "font_size_pt": rf.font_size_pt, "line_spacing_pt": p.format.line_spacing_pt, "bold": rf.bold, "color": getattr(rf, 'color', None)},
            "runs": runs_data,
        })

    result_tables = []
    for t in model.tables:
        cells = []
        for c in t.cells:
            cell_paras = [{"text": cp.text, "format": {"alignment": cp.format.alignment, "font_name": (cp.runs[0].format.font_name if cp.runs else None), "font_size_pt": (cp.runs[0].format.font_size_pt if cp.runs else None), "bold": (cp.runs[0].format.bold if cp.runs else None)}} for cp in c.paragraphs]
            cells.append({"row": c.row, "col": c.col, "text": c.text, "paragraphs": cell_paras})
        result_tables.append({"index": t.index, "rows": t.rows, "cols": t.cols, "cells": cells, "insert_after_index": t.insert_after_index})

    return {
        "success": True, "changes": changes,
        "paragraphs": result_paras, "tables": result_tables,
        "page_setup": page_setup, "doc_type": body.doc_type, "template_name": template_name,
    }


# ---------------------------------------------------------------------------
#  从预览数据生成 docx 并下载
# ---------------------------------------------------------------------------

class PreviewDownloadRequest(BaseModel):
    paragraphs: list[ParagraphData]
    tables: list[dict] | None = None
    page_setup: dict | None = None
    stamp: dict | None = None
    header_config: dict | None = None
    footer_note_config: dict | None = None
    page_number_config: dict | None = None
    source_doc_id: int | None = None  # 源文档 ID，用于保留原始文档结构


@router.post("/preview-download")
async def download_from_preview(body: PreviewDownloadRequest):
    """从前端预览数据（段落+表格）生成 docx 并返回下载。

    当 source_doc_id 有值时，使用原始 docx 作为基础（保留表格位置等结构）。
    否则从零创建（Markdown 转 docx 场景）。
    """
    from core.document.models import (
        DocumentModel, DocumentMetadata, PageSetup,
        Paragraph, ParagraphFormat, Run, RunFormat,
        Table, TableCell,
    )
    from core.document.generator import generate_docx
    from core.document.parser import parse_docx
    import tempfile

    # 构建 DocumentModel
    ps = PageSetup()
    if body.page_setup:
        ps.margin_top_mm = body.page_setup.get('margin_top_mm', 37)
        ps.margin_bottom_mm = body.page_setup.get('margin_bottom_mm', 35)
        ps.margin_left_mm = body.page_setup.get('margin_left_mm', 28)
        ps.margin_right_mm = body.page_setup.get('margin_right_mm', 26)

    # 优先使用源文档（保留原始结构：表格位置、页眉页脚等）
    model = None
    if body.source_doc_id:
        from db.database import SessionLocal
        db = SessionLocal()
        try:
            from services import document_service as svc
            doc = svc.get_document(db, body.source_doc_id)
            if doc and doc.file_path:
                source_path = doc.file_path
                # 也尝试使用 optimized_path
                if doc.optimized_path:
                    from pathlib import Path as _P
                    op = _P(doc.optimized_path)
                    if op.exists():
                        source_path = str(op)
                model = parse_docx(source_path)
                logger.info(f"Loaded source document {body.source_doc_id} for preview download")
        except Exception as e:
            logger.warning(f"Failed to load source document: {e}, creating new model")
        finally:
            db.close()

    if model is None:
        model = DocumentModel(
            metadata=DocumentMetadata(), page_setup=ps,
            paragraphs=[], tables=[], headers=[], footers=[],
        )
    else:
        # 用前端传来的页边距覆盖
        model.page_setup = ps

    # 用前端传来的段落替换模型段落
    model.paragraphs = []
    for i, p in enumerate(body.paragraphs):
        rf = RunFormat(font_name=p.format.get('font_name'), font_size_pt=p.format.get('font_size_pt'), bold=p.format.get('bold'))
        pf = ParagraphFormat(alignment=p.format.get('alignment'), first_line_indent_pt=p.format.get('first_line_indent_pt'), line_spacing_pt=p.format.get('line_spacing_pt'))
        model.paragraphs.append(Paragraph(index=i, text=p.text, is_heading=p.is_heading, heading_level=p.heading_level, role=p.role, runs=[Run(index=0, text=p.text, format=rf)], format=pf))

    # 还原表格
    model.tables = []
    if body.tables:
        for t_data in body.tables:
            table = Table(index=len(model.tables), rows=t_data.get('rows', 0), cols=t_data.get('cols', 0), cells=[])
            for c_data in t_data.get('cells', []):
                cell_paras = []
                for cp_data in c_data.get('paragraphs', []):
                    fmt = cp_data.get('format', {})
                    cp_rf = RunFormat(font_name=fmt.get('font_name'), font_size_pt=fmt.get('font_size_pt'), bold=fmt.get('bold'))
                    cp_pf = ParagraphFormat(alignment=fmt.get('alignment'))
                    cell_paras.append(Paragraph(index=0, text=cp_data.get('text', ''), runs=[Run(index=0, text=cp_data.get('text', ''), format=cp_rf)], format=cp_pf))
                table.cells.append(TableCell(row=c_data['row'], col=c_data['col'], text=c_data.get('text', ''), paragraphs=cell_paras))
            model.tables.append(table)

    # 生成 docx
    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False)
    tmp.close()
    output_path = generate_docx(model, tmp.name)

    # 版头/版记注入
    if body.header_config and body.header_config.get('enabled', True):
        _inject_header_to_docx(output_path, body.header_config)
    if body.footer_note_config and body.footer_note_config.get('enabled', True):
        _inject_footer_to_docx(output_path, body.footer_note_config)

    # 嵌入印章（如有）
    if body.stamp and body.stamp.get('image_base64'):
        _embed_stamp(
            output_path=str(output_path),
            image_base64=body.stamp['image_base64'],
            page_number=body.stamp.get('page_number', 0),
            x_mm=body.stamp.get('x_mm', 30),
            y_mm=body.stamp.get('y_mm', 25),
            width_mm=body.stamp.get('width_mm', 40),
            height_mm=body.stamp.get('height_mm', 40),
        )

    return FileResponse(
        path=str(output_path),
        filename="公文预览.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        background=BackgroundTask(os.unlink, str(output_path)),
    )


# ---------------------------------------------------------------------------
#  文档优化
# ---------------------------------------------------------------------------


@router.post("/{doc_id}", response_model=OptimizeResponse)
async def run_optimize(doc_id: int, req: OptimizeRequest | None = None, db: Session = Depends(get_db)):
    """Run auto-optimization on a document."""
    doc_type = req.document_type if req else None
    apply_fixes = req.apply_fixes if req else True
    selected_rule_ids = req.selected_rule_ids if req else None
    header_config = req.header_config if req else None
    footer_note_config = req.footer_note_config if req else None
    try:
        result = svc.optimize_document(
            db, doc_id, doc_type, apply_fixes, selected_rule_ids,
            header_config=header_config,
            footer_note_config=footer_note_config,
        )
    except ValueError as e:
        msg = str(e)
        # 区分"文档未找到"(404)和其他错误(422/500)
        if "not found" in msg.lower() or "未找到" in msg:
            raise HTTPException(status_code=404, detail=msg)
        raise HTTPException(status_code=422, detail=msg)
    except Exception as e:
        logger.error(f"Optimize failed for doc {doc_id}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail="文档优化失败，请稍后重试")

    return OptimizeResponse(
        document_id=result["document_id"],
        output_path=result["output_path"],
        fixes_applied=result["fixes_applied"],
        message=f"优化完成，已应用 {result['fixes_applied']} 项修复",
    )


@router.get("/{doc_id}/download")
async def download_optimized(doc_id: int, db: Session = Depends(get_db)):
    """Download the optimized document."""
    doc = svc.get_document(db, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    # 优先使用 DB 中存储的 optimized_path（由 optimize_document 写入）
    out_path = None
    if doc.optimized_path:
        p = Path(doc.optimized_path)
        if p.exists():
            out_path = p
        else:
            # 尝试按文件名在 OUTPUT_DIR 中查找
            from config import OUTPUT_DIR
            fallback = OUTPUT_DIR / p.name
            if fallback.exists():
                out_path = fallback
    if not out_path:
        # 回退：按命名规则拼接路径
        out_name = Path(doc.filename).stem + _OPTIMIZED_SUFFIX
        from config import OUTPUT_DIR
        out_path = OUTPUT_DIR / out_name
        if not out_path.exists():
            raise HTTPException(status_code=404, detail="Optimized file not found. Run optimize first.")

    return FileResponse(
        path=str(out_path),
        filename=out_path.name,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
