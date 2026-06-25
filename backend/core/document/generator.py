"""
Document generator: converts DocumentModel back into a .docx file.

核心设计变更 (v1.4.0-rc3):
- 基于源文档原地修改，不再创建空 Document()
- 正确保留表格、图片、嵌入对象、分节符等 DocumentModel 未建模的内容
- 所有字体设置必须经过 font_utils 统一入口
- 支持表格写入、页眉/页脚写入
"""
from __future__ import annotations
from pathlib import Path
from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from core.document.models import DocumentModel, Paragraph, Run, Table as TableModel, HeaderFooter
from core.document.font_utils import (
    set_run_font, set_paragraph_font, validate_document_fonts,
    TITLE_FONT, BODY_FONT, LATIN_FONT,
)
from utils.logger import logger


def generate_docx(model: DocumentModel, output_path: Path | str) -> Path:
    """
    Generate a .docx file from a DocumentModel.

    策略：加载源文档 → 替换段落内容（保留表格在原文位置）→ 更新表格 → 更新页眉页脚 → 保存。
    这样可以保留 DocumentModel 未建模的内容（图片、嵌入对象、分节符等）。

    Args:
        model: The document model to generate from
        output_path: Path where the .docx file should be saved

    Returns:
        Path to the generated file
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    logger.info(f"Generating document: {output_path}")

    # 决定源文档：优先使用源文件，若不可用则创建新文档
    source_path = model.source_path
    if source_path and Path(source_path).exists():
        try:
            doc = Document(str(source_path))
            logger.debug(f"Loaded source document: {source_path}")
        except Exception as e:
            logger.warning(f"Failed to load source doc, creating new: {e}")
            doc = Document()
    else:
        doc = Document()
        logger.debug("No source document, created new Document()")

    # 0. Apply document-level font defaults (prevents Word from using MS Gothic)
    _apply_document_defaults(doc)

    # 1. Apply page setup
    _apply_page_setup(doc, model)

    # 2. Replace paragraphs in-place (preserving table positions)
    _replace_paragraphs(doc, model)

    # 3. Update tables with model data
    _update_tables(doc, model)

    # 4. Update headers and footers
    _update_headers_footers(doc, model)

    # 5. Update metadata
    _update_metadata(doc, model)

    # 6. Post-generation font validation
    font_issues = validate_document_fonts(doc)
    if font_issues:
        logger.warning(f"Found {len(font_issues)} font issues after generation:")
        for fi in font_issues[:5]:
            logger.warning(f"  para={fi['paragraph']}, run={fi['run']}, "
                           f"attr={fi['attribute']}, font={fi['font_name']}, "
                           f"text='{fi['text']}'")

    # 7. Save
    doc.save(str(output_path))
    logger.info(f"Document saved: {output_path} (font issues: {len(font_issues)})")

    return output_path


# ---------------------------------------------------------------------------
#  Document Defaults & Page Setup
# ---------------------------------------------------------------------------

def _apply_document_defaults(doc: Document):
    """
    设置文档级别的默认字体。
    通过在 styles.xml 中设置 docDefaults，确保即使某个 run 没有显式设置字体，
    Word 也不会使用 MS Gothic 等替代字体回退。
    """
    try:
        styles_element = doc.styles.element
        doc_defaults = styles_element.find(qn('w:docDefaults'))
        if doc_defaults is None:
            doc_defaults = OxmlElement('w:docDefaults')
            styles_element.insert(0, doc_defaults)

        rPrDefault = doc_defaults.find(qn('w:rPrDefault'))
        if rPrDefault is None:
            rPrDefault = OxmlElement('w:rPrDefault')
            doc_defaults.append(rPrDefault)

        rPr = rPrDefault.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            rPrDefault.append(rPr)

        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)

        rFonts.set(qn('w:ascii'), LATIN_FONT)
        rFonts.set(qn('w:hAnsi'), LATIN_FONT)
        rFonts.set(qn('w:eastAsia'), BODY_FONT)
        rFonts.set(qn('w:cs'), BODY_FONT)

        logger.debug("Document default fonts applied: eastAsia=仿宋_GB2312, latin=Times New Roman")
    except Exception as e:
        logger.warning(f"Failed to set document default fonts: {e}")


def _apply_page_setup(doc: Document, model: DocumentModel):
    """Apply page setup to the document."""
    if not doc.sections:
        return
    section = doc.sections[0]
    ps = model.page_setup

    if ps.paper_width_mm is not None and 50 <= ps.paper_width_mm <= 1000:
        section.page_width = Mm(ps.paper_width_mm)
    if ps.paper_height_mm is not None and 50 <= ps.paper_height_mm <= 1000:
        section.page_height = Mm(ps.paper_height_mm)
    if ps.margin_top_mm is not None and 0 <= ps.margin_top_mm <= 100:
        section.top_margin = Mm(ps.margin_top_mm)
    if ps.margin_bottom_mm is not None and 0 <= ps.margin_bottom_mm <= 100:
        section.bottom_margin = Mm(ps.margin_bottom_mm)
    if ps.margin_left_mm is not None and 0 <= ps.margin_left_mm <= 100:
        section.left_margin = Mm(ps.margin_left_mm)
    if ps.margin_right_mm is not None and 0 <= ps.margin_right_mm <= 100:
        section.right_margin = Mm(ps.margin_right_mm)


# ---------------------------------------------------------------------------
#  Paragraph Replacement (in-place, preserving table positions)
# ---------------------------------------------------------------------------

def _replace_paragraphs(doc: Document, model: DocumentModel):
    """
    替换文档中的段落内容，同时保留表格在原始位置。

    策略：
    1. 找到 body 直接子元素中的 <w:p> 元素（排除表格内的段落）
    2. 按顺序替换为 model.paragraphs 的内容
    3. 表格 <w:tbl> 元素保持不动
    """
    body = doc.element.body
    p_tag = qn('w:p')
    # 只取 body 的直接子 <w:p>，排除表格单元格内的段落
    all_p_elements = [child for child in body if child.tag == p_tag]

    model_paras = model.paragraphs

    # 替换策略：逐个替换已有的段落，多余的追加，多余的原文段落清除
    for idx, para_model in enumerate(model_paras):
        if idx < len(all_p_elements):
            # 替换已有段落的内容
            _replace_paragraph_content(doc, all_p_elements[idx], para_model)
        else:
            # model 比原文多的段落，追加到 body 末尾
            new_para = doc.add_paragraph()
            _apply_paragraph_format(new_para, para_model)
            _add_runs_to_paragraph(new_para, para_model)

    # 清除多余的原文段落（model 中没有对应的）
    if len(all_p_elements) > len(model_paras):
        for idx in range(len(model_paras), len(all_p_elements)):
            try:
                body.remove(all_p_elements[idx])
            except Exception:
                pass  # 已被移除则跳过

    logger.debug(f"Replaced {min(len(model_paras), len(all_p_elements))} paragraphs, "
                 f"added {max(0, len(model_paras) - len(all_p_elements))}, "
                 f"removed {max(0, len(all_p_elements) - len(model_paras))}")


def _replace_paragraph_content(doc: Document, p_element, para_model: Paragraph):
    """
    替换一个 <w:p> 元素的内容（清除旧文本 runs，写入新 runs），保留段落属性和图片。
    """
    # 清除文本 runs，但保留含图片/绘图的 runs
    for child in list(p_element):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'r':
            # 检查 run 是否包含图片（w:drawing 或 w:pict）
            has_image = False
            for sub in child:
                sub_tag = sub.tag.split('}')[-1] if '}' in sub.tag else sub.tag
                if sub_tag in ('drawing', 'pict'):
                    has_image = True
                    break
            if not has_image:
                p_element.remove(child)
        elif tag == 'hyperlink':
            # 保留超链接（可能包含图片）
            pass

    # 更新段落属性 (w:pPr)
    _update_pPr(p_element, para_model)

    # 获取对应的 python-docx Paragraph 对象以使用 API 添加 run
    para_obj = _find_paragraph_object(doc, p_element)
    if para_obj is not None:
        _add_runs_to_paragraph(para_obj, para_model)
    else:
        # 回退：直接操作 XML 添加 runs
        _add_runs_via_xml(p_element, para_model)


def _update_pPr(p_element, para_model: Paragraph):
    """更新 <w:p> 元素的 <w:pPr> 段落属性。
    关键原则：model 有值才替换，None 保留原文档格式不删除。"""
    fmt = para_model.format

    # 获取或创建 pPr
    pPr = p_element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p_element.insert(0, pPr)

    # 对齐方式：仅当 model 有值时替换
    if fmt.alignment:
        jc = pPr.find(qn('w:jc'))
        if jc is not None:
            pPr.remove(jc)
        jc = OxmlElement('w:jc')
        alignment_map = {
            "left": "left", "center": "center",
            "right": "right", "justify": "both",
        }
        jc.set(qn('w:val'), alignment_map.get(fmt.alignment, "left"))
        pPr.append(jc)

    # 缩进：仅当 model 有值时替换，否则保留原文档缩进
    has_indent = (fmt.first_line_indent_pt is not None or
                  fmt.left_indent_pt is not None or
                  fmt.right_indent_pt is not None)
    if has_indent:
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            pPr.remove(ind)
        ind = OxmlElement('w:ind')
        if fmt.first_line_indent_pt is not None:
            ind.set(qn('w:firstLine'), str(int(fmt.first_line_indent_pt * 20)))
            chars = int(round(fmt.first_line_indent_pt / 16 * 100))
            if chars > 0:
                ind.set(qn('w:firstLineChars'), str(chars))
        if fmt.left_indent_pt is not None:
            ind.set(qn('w:left'), str(int(fmt.left_indent_pt * 20)))
        if fmt.right_indent_pt is not None:
            ind.set(qn('w:right'), str(int(fmt.right_indent_pt * 20)))
        pPr.append(ind)

    # 行距：仅当 model 有值时替换，否则保留原文档行距
    has_spacing = (fmt.line_spacing_pt is not None or
                   fmt.space_before_pt is not None or
                   fmt.space_after_pt is not None)
    if has_spacing:
        spacing = pPr.find(qn('w:spacing'))
        if spacing is not None:
            pPr.remove(spacing)
        spacing = OxmlElement('w:spacing')
        if fmt.line_spacing_pt is not None:
            spacing_pt = max(6, min(200, fmt.line_spacing_pt))
            rule = fmt.line_spacing_rule or "exact"
            if rule == "multiple":
                # 倍数行距：w:line 值为 240 分之一行（如 1.5x = 360）
                # line_spacing_pt 存储的是 pt 值，需要反算回倍数
                # 公文标准字号16pt，1倍行距=240（即 16pt * 15 = 240）
                # pt → 240ths: value = spacing_pt / 16 * 240
                line_val = int(round(spacing_pt / 16 * 240))
                spacing.set(qn('w:line'), str(line_val))
                spacing.set(qn('w:lineRule'), 'auto')
            elif rule == "atLeast":
                spacing.set(qn('w:line'), str(int(spacing_pt * 20)))
                spacing.set(qn('w:lineRule'), 'atLeast')
            else:
                # exact (默认，符合GB/T 9704标准)
                spacing.set(qn('w:line'), str(int(spacing_pt * 20)))
                spacing.set(qn('w:lineRule'), 'exact')
        if fmt.space_before_pt is not None:
            spacing.set(qn('w:before'), str(int(fmt.space_before_pt * 20)))
        if fmt.space_after_pt is not None:
            spacing.set(qn('w:after'), str(int(fmt.space_after_pt * 20)))
        pPr.append(spacing)


def _find_paragraph_object(doc: Document, p_element):
    """通过 XML 元素找到对应的 python-docx Paragraph 对象。"""
    for para in doc.paragraphs:
        if para._element is p_element:
            return para
    return None


def _add_runs_to_paragraph(para, para_model: Paragraph):
    """使用 python-docx API 向段落添加 runs。"""
    if para_model.runs:
        for run_model in para_model.runs:
            run = para.add_run(run_model.text)
            _apply_run_format(run, run_model)
    else:
        if para_model.text:
            run = para.add_run(para_model.text)
            set_run_font(run, BODY_FONT)


def _add_runs_via_xml(p_element, para_model: Paragraph):
    """直接通过 XML 添加 runs（当无法找到 python-docx Paragraph 对象时的回退方案）。"""
    if para_model.runs:
        for run_model in para_model.runs:
            r = OxmlElement('w:r')
            # 添加 run 格式属性
            rPr = OxmlElement('w:rPr')
            fmt = run_model.format
            if fmt.font_name:
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), fmt.font_name)
                rFonts.set(qn('w:hAnsi'), fmt.font_name)
                rFonts.set(qn('w:eastAsia'), fmt.font_name)
                rPr.append(rFonts)
            if fmt.font_size_pt:
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), str(int(fmt.font_size_pt * 2)))  # half-points
                rPr.append(sz)
            if fmt.bold:
                rPr.append(OxmlElement('w:b'))
            if fmt.italic:
                rPr.append(OxmlElement('w:i'))
            if len(rPr) > 0:
                r.append(rPr)
            t = OxmlElement('w:t')
            t.text = run_model.text
            t.set(qn('xml:space'), 'preserve')
            r.append(t)
            p_element.append(r)
    elif para_model.text:
        r = OxmlElement('w:r')
        t = OxmlElement('w:t')
        t.text = para_model.text
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        p_element.append(r)


# ---------------------------------------------------------------------------
#  Table Writing
# ---------------------------------------------------------------------------

def _update_tables(doc: Document, model: DocumentModel):
    """
    更新文档中的表格内容。
    如果源文档有表格，更新其单元格内容。
    如果源文档没有表格但 model 有，在末尾添加。
    """
    existing_tables = list(doc.tables)
    model_tables = model.tables

    for idx, table_model in enumerate(model_tables):
        if idx < len(existing_tables):
            # 更新已有表格的单元格内容
            _update_table_content(existing_tables[idx], table_model)
        else:
            # 添加新表格
            _add_table(doc, table_model)

    logger.debug(f"Updated {min(len(model_tables), len(existing_tables))} tables, "
                 f"added {max(0, len(model_tables) - len(existing_tables))}")


def _update_table_content(table, table_model: TableModel):
    """更新已有表格的单元格内容。"""
    for cell_model in table_model.cells:
        try:
            cell = table.cell(cell_model.row, cell_model.col)
            # 更新单元格中的段落内容
            if cell_model.paragraphs:
                for p_idx, para_model in enumerate(cell_model.paragraphs):
                    if p_idx < len(cell.paragraphs):
                        # 替换已有段落
                        para = cell.paragraphs[p_idx]
                        # 清除旧 runs
                        for run in list(para.runs):
                            run._element.getparent().remove(run._element)
                        # 添加新 runs
                        _add_runs_to_paragraph(para, para_model)
                        # 更新段落格式（缩进、行距、对齐等）
                        _update_pPr(para._element, para_model)
                    else:
                        # 添加新段落
                        para = cell.add_paragraph()
                        _add_runs_to_paragraph(para, para_model)
                        _apply_paragraph_format(para, para_model)
            elif cell_model.text:
                # 没有详细段落信息，直接设置文本
                if cell.paragraphs:
                    para = cell.paragraphs[0]
                    for run in list(para.runs):
                        run._element.getparent().remove(run._element)
                    run = para.add_run(cell_model.text)
                    set_run_font(run, BODY_FONT)
        except Exception as e:
            logger.warning(f"Failed to update table cell ({cell_model.row},{cell_model.col}): {e}")


def _add_table(doc: Document, table_model: TableModel):
    """在文档末尾添加一个新表格。"""
    try:
        rows = max(1, table_model.rows)
        cols = max(1, table_model.cols)
        table = doc.add_table(rows=rows, cols=cols)

        # 设置表格样式（带边框）
        table.style = 'Table Grid'

        for cell_model in table_model.cells:
            try:
                cell = table.cell(cell_model.row, cell_model.col)
                # 清除默认段落
                for para in cell.paragraphs:
                    for run in list(para.runs):
                        run._element.getparent().remove(run._element)

                if cell_model.paragraphs:
                    for p_idx, para_model in enumerate(cell_model.paragraphs):
                        if p_idx < len(cell.paragraphs):
                            para = cell.paragraphs[p_idx]
                            _add_runs_to_paragraph(para, para_model)
                            _update_pPr(para._element, para_model)
                        else:
                            para = cell.add_paragraph()
                            _add_runs_to_paragraph(para, para_model)
                            _apply_paragraph_format(para, para_model)
                elif cell_model.text:
                    if cell.paragraphs:
                        run = cell.paragraphs[0].add_run(cell_model.text)
                        set_run_font(run, BODY_FONT)
            except Exception as e:
                logger.warning(f"Failed to write table cell ({cell_model.row},{cell_model.col}): {e}")

        logger.debug(f"Added table: {rows}x{cols}")
    except Exception as e:
        logger.error(f"Failed to add table: {e}")


# ---------------------------------------------------------------------------
#  Headers & Footers
# ---------------------------------------------------------------------------

def _update_headers_footers(doc: Document, model: DocumentModel):
    """更新页眉和页脚内容。"""
    _update_hf_list(doc, model.headers, "header")
    _update_hf_list(doc, model.footers, "footer")


def _update_hf_list(doc: Document, hf_models: list[HeaderFooter], hf_type: str):
    """更新一组页眉或页脚。"""
    for hf_model in hf_models:
        try:
            sec_idx = hf_model.section_index
            if sec_idx >= len(doc.sections):
                logger.debug(f"Section {sec_idx} not found, skipping {hf_type}")
                continue

            section = doc.sections[sec_idx]
            target = section.header if hf_type == "header" else section.footer

            if target is None:
                continue

            # 更新页眉/页脚中的段落
            if hf_model.paragraphs:
                for p_idx, para_model in enumerate(hf_model.paragraphs):
                    if p_idx < len(target.paragraphs):
                        para = target.paragraphs[p_idx]
                        # 清除旧 runs
                        for run in list(para.runs):
                            run._element.getparent().remove(run._element)
                        # 添加新 runs
                        _add_runs_to_paragraph(para, para_model)
                        # 更新段落格式（缩进、行距、对齐等）
                        _update_pPr(para._element, para_model)
                    else:
                        para = target.add_paragraph()
                        _add_runs_to_paragraph(para, para_model)
                        _apply_paragraph_format(para, para_model)
            elif hf_model.text:
                # 简单文本模式
                if target.paragraphs:
                    para = target.paragraphs[0]
                    for run in list(para.runs):
                        run._element.getparent().remove(run._element)
                    run = para.add_run(hf_model.text)
                    set_run_font(run, BODY_FONT)

            logger.debug(f"Updated {hf_type} section {sec_idx}")
        except Exception as e:
            logger.warning(f"Failed to update {hf_type} section {hf_model.section_index}: {e}")


# ---------------------------------------------------------------------------
#  Metadata
# ---------------------------------------------------------------------------

def _update_metadata(doc: Document, model: DocumentModel):
    """更新文档核心属性（元数据）。"""
    try:
        meta = model.metadata
        props = doc.core_properties
        if meta.title:
            props.title = meta.title
        if meta.author:
            props.author = meta.author
        if meta.subject:
            props.subject = meta.subject
        if meta.category:
            props.category = meta.category
        logger.debug("Document metadata updated")
    except Exception as e:
        logger.warning(f"Failed to update metadata: {e}")


# ---------------------------------------------------------------------------
#  Format Helpers
# ---------------------------------------------------------------------------

def _apply_paragraph_format(para, para_model: Paragraph):
    """Apply formatting to a paragraph using python-docx API."""
    pf = para.paragraph_format
    fmt = para_model.format

    # Alignment
    if fmt.alignment:
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }
        para.alignment = alignment_map.get(fmt.alignment, WD_ALIGN_PARAGRAPH.LEFT)

    # Indentation
    if fmt.first_line_indent_pt is not None:
        pf.first_line_indent = Pt(fmt.first_line_indent_pt)
    if fmt.left_indent_pt is not None:
        pf.left_indent = Pt(fmt.left_indent_pt)
    if fmt.right_indent_pt is not None:
        pf.right_indent = Pt(fmt.right_indent_pt)

    # Spacing
    if fmt.space_before_pt is not None:
        pf.space_before = Pt(fmt.space_before_pt)
    if fmt.space_after_pt is not None:
        pf.space_after = Pt(fmt.space_after_pt)

    # Line spacing
    if fmt.line_spacing_pt is not None:
        spacing_pt = max(6, min(200, fmt.line_spacing_pt))
        pf.line_spacing = Pt(spacing_pt)
        rule = fmt.line_spacing_rule or "exact"
        if rule == "multiple":
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        elif rule == "atLeast":
            pf.line_spacing_rule = WD_LINE_SPACING.AT_LEAST
        else:
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY


def _apply_run_format(run, run_model: Run):
    """
    Apply formatting to a run.
    使用 font_utils.set_run_font 统一处理中文字体。
    """
    fmt = run_model.format

    # === 字体设置（统一入口） ===
    if fmt.font_name:
        set_run_font(run, fmt.font_name)
    else:
        set_run_font(run, BODY_FONT)

    # === 字号 ===
    if fmt.font_size_pt is not None:
        run.font.size = Pt(fmt.font_size_pt)

    # === 样式 ===
    if fmt.bold is not None:
        run.font.bold = fmt.bold
    if fmt.italic is not None:
        run.font.italic = fmt.italic
    if fmt.underline is not None:
        run.font.underline = fmt.underline

    # === 颜色 ===
    if fmt.color:
        try:
            rgb_str = fmt.color.replace("#", "")
            if len(rgb_str) == 6:
                r = int(rgb_str[0:2], 16)
                g = int(rgb_str[2:4], 16)
                b = int(rgb_str[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
        except Exception:
            pass
