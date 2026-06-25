# This file is part of the Official Document AI Assistant.
# (c) 2026 Jose AI (https://www.linhut.cn)
# Licensed under the MIT License. See the LICENSE file for details.
"""
Font utilities: Handle Chinese font settings correctly for Word documents.
解决中文字体在Word中显示为MS Gothic等问题

核心原则：
- 所有字体设置必须经过此模块的统一入口
- 禁止项目任何位置直接调用 run.font.name
- 必须同时设置 w:rFonts 的 ascii, hAnsi, eastAsia, cs 四个属性

参考：
- GB/T 9704-2012 党政机关公文格式
- python-docx 的 OXML 底层字体机制
- AIPoliDoc/MCP-Doc 项目的文档处理思想
"""
from __future__ import annotations
from typing import Optional
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from utils.logger import logger


# ---------------------------------------------------------------------------
#  标准字体定义
# ---------------------------------------------------------------------------

# 公文标题字体
TITLE_FONT = "方正小标宋简体"
TITLE_FONT_FALLBACK = "SimSun"

# 公文正文字体
BODY_FONT = "仿宋_GB2312"
BODY_FONT_FALLBACK = "FangSong"

# 公文西文/数字字体（固定）
LATIN_FONT = "Times New Roman"

# 字体回退映射表
FONT_FALLBACK_MAP = {
    "方正小标宋简体": "SimSun",
    "仿宋_GB2312": "FangSong",
    "楷体_GB2312": "KaiTi",
    "黑体": "SimHei",
    "宋体": "SimSun",
    "FangSong": "仿宋",
    "KaiTi": "楷体",
    "SimHei": "黑体",
    "SimSun": "宋体",
}

# 无效字体名字模式（这些出现在文档中就说明字体设置有问题）
INVALID_FONT_PATTERNS = [
    "MS Gothic",
    "MS Mincho",
    "MS PGothic",
    "MS PMincho",
    "ＭＳ ゴシック",
    "ＭＳ 明朝",
]


def set_run_font(run, font_name: str, latin_font: str | None = None) -> None:
    """
    为 run 设置字体，正确处理中文字体。

    Word 文档需要同时设置四个字体属性：
    1. w:ascii   - 西文 ASCII 字体（0-127）
    2. w:hAnsi   - 高位 ANSI 字体（128-255，含西文扩展）
    3. w:eastAsia - 东亚字体（中文、日文、韩文）
    4. w:cs      - 复杂脚本字体（阿拉伯文、泰文等）

    只设置 run.font.name 只会写入 w:ascii 和 w:hAnsi，
    不会写入 w:eastAsia，导致 Word 使用默认东亚字体（MS Gothic）。

    Args:
        run: python-docx Run 对象
        font_name: 中文字体名称
        latin_font: 西文/数字字体（默认使用 Times New Roman）
    """
    if not font_name:
        return

    latin = latin_font or LATIN_FONT

    # 1. 通过 python-docx API 设置基础字体（这会设置 w:ascii + w:hAnsi）
    run.font.name = latin

    # 2. 获取或创建 rPr 元素
    rPr = run._element.get_or_add_rPr()

    # 3. 获取或创建 rFonts 元素
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)

    # 4. 设置四个关键字体属性
    rFonts.set(qn('w:ascii'), latin)       # 西文 ASCII
    rFonts.set(qn('w:hAnsi'), latin)       # 高位 ANSI
    rFonts.set(qn('w:eastAsia'), font_name)  # 东亚文字（核心！）
    rFonts.set(qn('w:cs'), font_name)       # 复杂脚本

    logger.debug(f"Set font: eastAsia={font_name}, latin={latin}")


def set_run_font_east_asian(run, font_name: str) -> None:
    """
    设置 run 的东亚字体（中文）。
    保留西文字体不变。

    Args:
        run: python-docx Run 对象
        font_name: 东亚字体名称
    """
    rPr = run._element.get_or_add_rPr()

    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)

    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def set_paragraph_font(paragraph, font_name: str, latin_font: str | None = None) -> None:
    """
    为段落中所有 run 设置字体。

    Args:
        paragraph: python-docx Paragraph 对象
        font_name: 中文字体名称
        latin_font: 西文字体（默认 Times New Roman）
    """
    for run in paragraph.runs:
        set_run_font(run, font_name, latin_font)


def apply_paragraph_style_font(paragraph, font_name: str, font_size_pt: float,
                                bold: bool = False, latin_font: str | None = None) -> None:
    """
    为段落中的所有 run 应用统一的字体格式。
    这是一个便捷的综合设置函数。

    Args:
        paragraph: python-docx Paragraph 对象
        font_name: 中文字体名称
        font_size_pt: 字号（磅）
        bold: 是否加粗
        latin_font: 西文字体
    """
    # 如果已有 runs，设置字体
    if paragraph.runs:
        for run in paragraph.runs:
            set_run_font(run, font_name, latin_font)
            run.font.size = Pt(font_size_pt)
            run.font.bold = bold
    else:
        # 没有 runs 则添加一个新的
        run = paragraph.add_run(paragraph.text)
        set_run_font(run, font_name, latin_font)
        run.font.size = Pt(font_size_pt)
        run.font.bold = bold


def detect_font_from_run(run) -> dict[str, str | None]:
    """
    从 run 的 XML 中读取所有字体属性。
    用于解析已有文档中的实际字体设置。

    Returns:
        dict with keys: ascii, hAnsi, eastAsia, cs, font_name
    """
    result = {
        "ascii": None,
        "hAnsi": None,
        "eastAsia": None,
        "cs": None,
        "font_name": None,
    }

    # 从 python-docx API 读取
    if run.font.name:
        result["font_name"] = run.font.name

    # 从底层 XML 读取
    rPr = run._element.find(qn('w:rPr'))
    if rPr is not None:
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is not None:
            result["ascii"] = rFonts.get(qn('w:ascii'))
            result["hAnsi"] = rFonts.get(qn('w:hAnsi'))
            result["eastAsia"] = rFonts.get(qn('w:eastAsia'))
            result["cs"] = rFonts.get(qn('w:cs'))

    return result


def get_effective_font(run) -> str | None:
    """
    获取 run 的有效字体。
    优先级：eastAsia > font.name

    这是解析文档时应该使用的函数，因为 Word 用 eastAsia 渲染中文。
    """
    font_info = detect_font_from_run(run)

    # eastAsia 是实际渲染中文的字体
    if font_info["eastAsia"]:
        return font_info["eastAsia"]

    # 回退到 font.name
    return font_info["font_name"]


def get_font_fallback(font_name: str) -> str:
    """
    获取字体的回退版本。
    当指定字体在系统中不可用时使用。

    Args:
        font_name: 目标字体名称

    Returns:
        回退字体名称
    """
    return FONT_FALLBACK_MAP.get(font_name, font_name)


def validate_font_name(font_name: str | None) -> bool:
    """
    验证字体名称是否合理。
    检测是否出现了无效的替代字体（如 MS Gothic）。

    Args:
        font_name: 字体名称

    Returns:
        True 如果字体有效
    """
    if not font_name:
        return False

    for pattern in INVALID_FONT_PATTERNS:
        if pattern.lower() in font_name.lower():
            return False

    return True


def validate_document_fonts(doc) -> list[dict]:
    """
    检查整个文档中是否存在无效字体（含表格、页眉页脚）。
    用于生成后的质量检查和自动修复。

    Args:
        doc: python-docx Document 对象

    Returns:
        所有发现的无效字体报告列表（含 run_obj 引用用于自动修复）
    """
    issues = []

    def _check_runs(runs, location_prefix):
        for run_idx, run in enumerate(runs):
            font_info = detect_font_from_run(run)
            for attr in ["ascii", "hAnsi", "eastAsia", "cs"]:
                fname = font_info.get(attr)
                if fname and not validate_font_name(fname):
                    issues.append({
                        "paragraph": location_prefix,
                        "run": run_idx,
                        "attribute": attr,
                        "font_name": fname,
                        "text": run.text[:50],
                        "run_obj": run,  # 保留引用用于自动修复
                    })

    # 1. 正文段落
    for para_idx, para in enumerate(doc.paragraphs):
        _check_runs(para.runs, f"body:{para_idx}")

    # 2. 表格单元格
    for tbl_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    _check_runs(para.runs, f"table:{tbl_idx}.{row_idx}.{col_idx}")

    # 3. 页眉页脚
    for section_idx, section in enumerate(doc.sections):
        if section.header:
            for para in section.header.paragraphs:
                _check_runs(para.runs, f"header:{section_idx}")
        if section.footer:
            for para in section.footer.paragraphs:
                _check_runs(para.runs, f"footer:{section_idx}")

    return issues