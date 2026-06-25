# This file is part of the Official Document AI Assistant.
# (c) 2026 Jose AI (https://www.linhut.cn)
# Licensed under the MIT License. See the LICENSE file for details.
"""
Document parser: converts a .docx file into DocumentModel.
Uses python-docx to read the file structure.

设计原则：
- 完整保留所有段落、Run、样式、字体、间距等格式信息
- 解析过程中禁止丢失任何格式数据
- 使用 font_utils.detect_font_from_run 正确获取中文字体（eastAsia）
"""
from __future__ import annotations
from pathlib import Path
from typing import Optional
import re

from docx import Document
from docx.shared import Pt, Cm, Inches, Length
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from core.document.models import (
    DocumentModel, DocumentMetadata, Paragraph, Run, RunFormat,
    ParagraphFormat, PageSetup, Table, TableCell, HeaderFooter
)
from core.document.font_utils import detect_font_from_run, get_effective_font
from utils.logger import logger

# 中文公文标题特征字体映射 → heading_level
_HEADING_FONT_MAP = {
    "方正小标宋简体": 0,  # 公文大标题
    "小标宋": 0,
    "黑体": 1,            # 一级标题
    "SimHei": 1,
    "楷体": 2,            # 二级标题
    "楷体_gb2312": 2,
    "KaiTi": 2,
}

# 一级标题序号模式（"一、""二、"等）
_H1_PATTERN = re.compile(r'^[一二三四五六七八九十]+、')
# 二级标题序号模式（"（一）""（二）"等）
_H2_PATTERN = re.compile(r'^（[一二三四五六七八九十]+）')
# 三级标题序号模式（"1.""2."等阿拉伯数字+句点）
_H3_PATTERN = re.compile(r'^\d+[.]')
# 四级标题序号模式（"（1）""（2）"等半角/全角括号+阿拉伯数字）
_H4_PATTERN = re.compile(r'^[（(]\d+[）)]')


def parse_docx(file_path: Path | str) -> DocumentModel:
    """
    Parse a .docx file into a DocumentModel.

    Args:
        file_path: Path to the .docx file

    Returns:
        DocumentModel instance with full format preservation
    """
    file_path = Path(file_path)
    doc = Document(str(file_path))

    logger.info(f"Parsing document: {file_path}")

    # 1. 解析元数据
    metadata = _parse_metadata(doc)

    # 2. 解析页面设置
    page_setup = _parse_page_setup(doc)

    # 3. 解析页眉页脚
    headers = _parse_headers_footers(doc, "header")
    footers = _parse_headers_footers(doc, "footer")

    # 4. 解析段落
    paragraphs = []
    for idx, para in enumerate(doc.paragraphs):
        parsed_para = _parse_paragraph(para, idx)
        paragraphs.append(parsed_para)

    # 4.1 后处理：位置/统计启发式（补偿未排版文档的标题检测不足）
    _post_detect_headings(paragraphs)

    # 4.2 段落角色标注
    _assign_paragraph_roles(paragraphs)

    # 5. 解析表格
    tables = []
    for idx, table in enumerate(doc.tables):
        parsed_table = _parse_table(table, idx)
        tables.append(parsed_table)

    model = DocumentModel(
        metadata=metadata,
        page_setup=page_setup,
        paragraphs=paragraphs,
        tables=tables,
        headers=headers,
        footers=footers,
        source_path=str(file_path),
        filename=file_path.name,
    )

    logger.info(f"Parsed: {len(paragraphs)} paragraphs, {len(tables)} tables, "
                f"{len(headers)} headers, {len(footers)} footers")

    # 6. AI辅助结构分析（当启发式检测不足时尝试）
    try:
        from core.document.ai_structure_analyzer import should_use_ai_analysis, classify_with_ai
        if should_use_ai_analysis(model):
            logger.info("Heading detection insufficient, attempting AI structure analysis...")
            if classify_with_ai(model):
                headings = [p for p in model.paragraphs if p.is_heading]
                logger.info(f"AI analysis complete, total headings now: {len(headings)}")
    except Exception as e:
        logger.debug(f"AI structure analysis skipped: {e}")

    return model


# ---------------------------------------------------------------------------
#  Metadata
# ---------------------------------------------------------------------------

def _parse_metadata(doc: Document) -> DocumentMetadata:
    """Parse document metadata (core properties)."""
    props = doc.core_properties
    return DocumentMetadata(
        title=props.title,
        author=props.author,
        subject=props.subject,
        created=str(props.created) if props.created else None,
        modified=str(props.modified) if props.modified else None,
        category=props.category,
    )


# ---------------------------------------------------------------------------
#  Page Setup
# ---------------------------------------------------------------------------

def _parse_page_setup(doc: Document) -> PageSetup:
    """Parse full page setup from document section."""
    if not doc.sections:
        return PageSetup()

    section = doc.sections[0]

    return PageSetup(
        paper_width_mm=_safe_mm(section.page_width, 210),
        paper_height_mm=_safe_mm(section.page_height, 297),
        margin_top_mm=_safe_mm(section.top_margin, 37),
        margin_bottom_mm=_safe_mm(section.bottom_margin, 35),
        margin_left_mm=_safe_mm(section.left_margin, 28),
        margin_right_mm=_safe_mm(section.right_margin, 26),
        orientation="landscape" if section.orientation == 1 else "portrait",
    )


def _safe_mm(value, default: float) -> float:
    """Safely convert a docx length to millimeters."""
    try:
        if value is not None:
            mm = value.mm
            if 10 <= mm <= 1000:
                return round(mm, 2)
    except Exception:
        pass
    return default


# ---------------------------------------------------------------------------
#  Headers / Footers
# ---------------------------------------------------------------------------

def _parse_headers_footers(doc: Document, hf_type: str) -> list[HeaderFooter]:
    """Parse headers or footers from all sections, including page number detection."""
    result = []
    for sec_idx, section in enumerate(doc.sections):
        target = section.header if hf_type == "header" else section.footer
        if target is None:
            continue

        # Build text from all paragraphs
        texts = []
        paras = []
        has_page_num = False
        for p_idx, para in enumerate(target.paragraphs):
            if para.text.strip():
                texts.append(para.text)
            paras.append(_parse_paragraph(para, p_idx))
            # 检测页码域代码（通过XML层）
            if _paragraph_has_page_field(para):
                has_page_num = True

        result.append(HeaderFooter(
            section_index=sec_idx,
            type=hf_type,
            text="\n".join(texts),
            paragraphs=paras,
            has_page_number=has_page_num,
        ))

    return result


def _paragraph_has_page_field(para) -> bool:
    """
    检测段落是否包含 Word 页码域代码（PAGE field）。
    通过解析 XML 层的 w:fldSimple 或 w:fldChar/w:instrText 元素判断。
    """
    try:
        from lxml import etree
        nsmap = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        xml = para._element

        # 方式1: 简单域 <w:fldSimple w:instr=" PAGE ">
        for fld_simple in xml.findall('.//w:fldSimple', nsmap):
            instr = fld_simple.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}instr', '')
            if 'PAGE' in instr.upper():
                return True

        # 方式2: 复杂域 <w:instrText> PAGE </w:instrText>
        for instr_text in xml.findall('.//w:instrText', nsmap):
            if instr_text.text and 'PAGE' in instr_text.text.upper():
                return True

    except Exception:
        pass
    return False


# ---------------------------------------------------------------------------
#  Paragraph
# ---------------------------------------------------------------------------

def _parse_paragraph(para, index: int) -> Paragraph:
    """Parse a single paragraph with full format preservation."""
    # 样式信息
    style_name = para.style.name if para.style else None
    is_heading = False
    heading_level = None

    if style_name:
        style_lower = style_name.lower()
        if style_lower.startswith('heading') or style_name.startswith('标题') or style_name.startswith('Heading'):
            is_heading = True
            # 提取层级：只在确认是 heading 样式时才提取数字
            for i in range(9, 0, -1):
                if f'heading {i}' in style_lower or f'heading{i}' in style_lower:
                    heading_level = i
                    break
            if heading_level is None:
                heading_level = 1

    # 段落格式
    para_format = _parse_paragraph_format(para)

    # Run 格式
    runs = []
    for run_idx, run in enumerate(para.runs):
        parsed_run = _parse_run(run, run_idx)
        runs.append(parsed_run)

    # 如果没有 run 但有文本，创建一个文本 run
    if not runs and para.text:
        runs.append(Run(
            text=para.text,
            index=0,
            format=RunFormat(),
        ))

    # ---- 段落样式 fallback：当 run 没有直接格式时，从段落样式读取 ----
    _apply_style_fallback(para, runs, para_format)

    # ---- 内容验证：Word 样式可能是错的 ----
    # 很多文档模板将正文误标为 Heading 3，需要验证内容是否真的像标题
    if is_heading and heading_level and heading_level >= 2:
        text_stripped = para.text.strip() if para.text else ""
        # 检查是否有标题内容信号
        has_content_signal = False
        if text_stripped:
            has_content_signal = bool(
                _H1_PATTERN.match(text_stripped) or
                _H2_PATTERN.match(text_stripped) or
                _H3_PATTERN.match(text_stripped)
            )
        # 检查是否有标题格式信号（加粗或标题字体）
        has_format_signal = False
        for run in runs:
            if run.text.strip():
                font = (run.format.font_name or "").lower()
                if run.format.bold or "黑体" in font or "楷体" in font:
                    has_format_signal = True
                break
        # 既无内容信号也无格式信号 → 降级为正文
        if not has_content_signal and not has_format_signal:
            is_heading = False
            heading_level = None

    # ---- 中文公文启发式标题检测 ----
    # 无论 Word 样式是否识别为标题，都运行启发式检测。
    # 因为 Word 可能将"方正小标宋简体"标题误判为 Heading 1，
    # 而启发式可以将其正确识别为 heading_level=0（公文大标题）。
    if runs:
        heuristic_heading, heuristic_level = _detect_heading_heuristic(para.text, runs, para_format)
        if heuristic_heading:
            if not is_heading:
                # Word 样式未识别为标题，使用启发式结果
                is_heading = True
                heading_level = heuristic_level
            elif heuristic_level == 0 and heading_level != 0:
                # Word 样式识别为 heading_N，但启发式识别为公文大标题（level 0）
                # 公文大标题（方正小标宋简体）优先级更高
                heading_level = 0

    return Paragraph(
        text=para.text,
        index=index,
        style_name=style_name,
        is_heading=is_heading,
        heading_level=heading_level,
        format=para_format,
        runs=runs,
    )


def _detect_heading_heuristic(
    text: str, runs: list[Run], para_format: ParagraphFormat
) -> tuple[bool, int | None]:
    """
    中文公文标题启发式检测（v2: 支持未排版文档）。

    检测策略：
    1. 格式信号：字体名称、字号、加粗、对齐（准确率高）
    2. 内容信号：编号模式、文本长度、位置（无需格式信息）

    两种信号独立工作，任一命中即可识别标题。
    未排版文档主要依赖内容信号。

    返回: (is_heading, heading_level)
    """
    text_stripped = text.strip() if text else ""
    if not text_stripped or len(text_stripped) > 80:
        return False, None  # 空文本或超长文本不作为标题

    # 获取主run的字体信息（取第一个非空run）
    main_font = None
    main_size = None
    is_bold = False
    for run in runs:
        if run.text.strip():
            main_font = (run.format.font_name or "").strip()
            main_size = run.format.font_size_pt
            is_bold = run.format.bold or False
            break

    font_lower = (main_font or "").lower()
    alignment = para_format.alignment or "left"
    has_font_signal = bool(main_font)  # 是否有有效字体信息

    # =====================================================================
    #  路径A：格式信号驱动（高置信度，已排版文档走这条路径）
    # =====================================================================

    # --- Level 0: 公文大标题 ---
    if has_font_signal:
        if "小标宋" in font_lower or main_font in ("方正小标宋简体",):
            return True, 0
        if main_size and main_size >= 20 and alignment == "center":
            return True, 0

    # --- Level 1: 一级标题（黑体）---
    if has_font_signal and ("黑体" in font_lower or font_lower in ("simhei",)):
        if alignment == "center" or is_bold or len(text_stripped) < 30:
            return True, 1

    # 格式信号："一、" + 加粗或黑体
    if _H1_PATTERN.match(text_stripped) and len(text_stripped) < 50:
        if is_bold or "黑体" in font_lower:
            return True, 1

    # --- Level 2: 二级标题（楷体 + 加粗）---
    if has_font_signal and ("楷体" in font_lower or font_lower in ("kaiti", "楷体_gb2312")) and is_bold:
        return True, 2

    # "（一）" 格式（无论字体如何，此模式足够唯一）
    if _H2_PATTERN.match(text_stripped) and len(text_stripped) < 50:
        return True, 2

    # --- Level 3: 三级标题（仿宋加粗 或 "1." + 加粗）---
    if has_font_signal and ("仿宋" in font_lower or font_lower in ("fangsong", "仿宋_gb2312")) and is_bold:
        if len(text_stripped) < 50:
            return True, 3

    if _H3_PATTERN.match(text_stripped) and is_bold and len(text_stripped) < 60:
        return True, 3

    # =====================================================================
    #  路径B：内容信号驱动（未排版文档走这条路径）
    #  当格式信号不足时，仅靠编号模式+文本长度判断
    # =====================================================================

    # B-0: 居中短文本 → 很可能是公文标题（即使无特殊字体）
    if alignment == "center" and len(text_stripped) < 30 and not _H1_PATTERN.match(text_stripped):
        return True, 0

    # B-1: "一、" 开头 + 短文本 → 一级标题（无需加粗/黑体）
    if _H1_PATTERN.match(text_stripped) and len(text_stripped) < 40:
        return True, 1

    # B-3: "1." 开头 + 短文本 → 三级标题（无需加粗）
    if _H3_PATTERN.match(text_stripped) and len(text_stripped) < 40:
        return True, 3

    # B-4: "（1）" 开头 + 短文本 → 四级标题
    if _H4_PATTERN.match(text_stripped) and len(text_stripped) < 40:
        return True, 4

    return False, None


def _post_detect_headings(paragraphs: list[Paragraph]) -> None:
    """
    后处理：位置/统计启发式标题检测。

    解决未排版文档中标题检测不足的问题：
    1. 若无 level=0 标题，将第一个短段落标记为标题
    2. 统计段落平均长度，短于平均40%的编号段落标记为标题候选
    """
    non_empty = [p for p in paragraphs if p.text.strip()]
    if len(non_empty) < 3:
        return

    # --- 补检公文标题 (level 0) ---
    has_title = any(p.is_heading and p.heading_level == 0 for p in paragraphs)
    if not has_title:
        # 第一个非空段落：若短于30字符且无句末标点，视为标题
        first = non_empty[0]
        t = first.text.strip()
        if len(t) < 30 and not t.endswith(('。', '；', '，', '！', '？', '.', ';', ',')):
            first.is_heading = True
            first.heading_level = 0
            logger.info(f"[后处理] 将首段识别为公文标题: {t[:30]!r}")

    # --- 统计辅助：基于段落长度的标题候选 ---
    lengths = [len(p.text.strip()) for p in non_empty if len(p.text.strip()) > 0]
    if not lengths:
        return
    avg_len = sum(lengths) / len(lengths)

    # 已检测到的标题数
    detected_count = sum(1 for p in paragraphs if p.is_heading)

    # 若标题数过少（< 2个），尝试用长度+编号模式补充检测
    if detected_count < 2 and avg_len > 20:
        for p in paragraphs:
            if p.is_heading or not p.text.strip():
                continue
            t = p.text.strip()
            # 短于平均长度40% + 匹配编号模式 → 标题候选
            if len(t) < avg_len * 0.4:
                if _H1_PATTERN.match(t) and len(t) < 40:
                    p.is_heading = True
                    p.heading_level = 1
                    logger.info(f"[后处理] 统计+模式补充检测一级标题: {t[:30]!r}")
                elif _H2_PATTERN.match(t) and len(t) < 40:
                    p.is_heading = True
                    p.heading_level = 2
                    logger.info(f"[后处理] 统计+模式补充检测二级标题: {t[:30]!r}")
                elif _H3_PATTERN.match(t) and len(t) < 40:
                    p.is_heading = True
                    p.heading_level = 3
                    logger.info(f"[后处理] 统计+模式补充检测三级标题: {t[:30]!r}")


# ---------------------------------------------------------------------------
#  Paragraph Role Assignment (段落角色标注)
# ---------------------------------------------------------------------------

# 日期模式
_DATE_RE = re.compile(r'^\d{4}年\d{1,2}月\d{1,2}日$')
_DATE_ALT_RE = re.compile(r'^\d{4}[.\-/]\d{1,2}[.\-/]\d{1,2}$')

# 主送机关特征词
_RECIPIENT_KEYWORDS = [
    '各部门', '各科室', '各单位', '全体', '各有关', '各相关',
    '局', '办', '委', '厅', '处', '室', '院', '中心',
]

# 附件模式
_ATTACHMENT_RE = re.compile(r'^附件[：:]')

# 抄送模式
_CC_RE = re.compile(r'^抄送[：:]')

# 印发模式
_PRINT_RE = re.compile(r'^印发机关|^印发日期')


def _assign_paragraph_roles(paragraphs: list[Paragraph]) -> None:
    """
    后处理：为每个段落标注 role（角色）。
    role 值: title / recipient / body / signature / date / attachment / cc / notes
    """
    non_empty = [(i, p) for i, p in enumerate(paragraphs) if p.text.strip()]
    if not non_empty:
        return

    # 第一段非空段落 = 标题
    first_idx, first_para = non_empty[0]
    if first_para.is_heading and first_para.heading_level == 0:
        first_para.role = 'title'
    elif not first_para.is_heading and len(first_para.text.strip()) < 30:
        first_para.role = 'title'

    # 最后几段：日期、落款、抄送、印发
    for i in range(len(non_empty) - 1, max(len(non_empty) - 6, 0), -1):
        idx, para = non_empty[i]
        text = para.text.strip()

        # 日期
        if _DATE_RE.match(text) or _DATE_ALT_RE.match(text):
            para.role = 'date'
            continue

        # 抄送
        if _CC_RE.match(text):
            para.role = 'cc'
            continue

        # 印发
        if _PRINT_RE.match(text):
            para.role = 'cc'
            continue

    # 落款：日期前的短文本（< 20 字，或含机关关键词）
    date_indices = [i for i, p in non_empty if p.role == 'date']
    if date_indices:
        date_idx = date_indices[-1]
        for i in range(len(non_empty)):
            idx, para = non_empty[i]
            if idx == date_idx - 1 and not para.is_heading:
                text = para.text.strip()
                if len(text) < 20:
                    para.role = 'signature'
                elif any(kw in text for kw in ['人民政府', '委员会', '办公厅', '办公室', '管理局', '局', '部']):
                    if len(text) < 40:
                        para.role = 'signature'

    # 主送机关：标题后第一段，以冒号结尾
    if len(non_empty) >= 2:
        second_idx, second_para = non_empty[1]
        if not second_para.role:
            text = second_para.text.strip()
            if text.endswith(('：', ':')) and len(text) < 50:
                second_para.role = 'recipient'
            elif any(kw in text for kw in _RECIPIENT_KEYWORDS) and text.endswith(('：', ':')):
                second_para.role = 'recipient'

    # 附件
    for idx, para in non_empty:
        if not para.role and _ATTACHMENT_RE.match(para.text.strip()):
            para.role = 'attachment'

    # 其余非空段落默认为 body
    for idx, para in non_empty:
        if not para.role:
            para.role = 'body'


# ---------------------------------------------------------------------------
#  Paragraph Format
# ---------------------------------------------------------------------------

def _apply_style_fallback(para, runs: list[Run], para_format: ParagraphFormat) -> None:
    """
    当 run 没有直接格式时，从段落样式读取默认值。
    确保预览时 font_size_pt / font_name 等字段不会是 None。
    """
    try:
        style = para.style
        if style is None:
            return

        # 读取样式的字体大小
        style_font_size = None
        try:
            if style.font and style.font.size:
                style_font_size = round(style.font.size.pt, 1)
        except Exception:
            pass

        # 读取样式的字体名称
        style_font_name = None
        try:
            if style.font and style.font.name:
                style_font_name = style.font.name
        except Exception:
            pass

        # 读取样式的行距
        style_line_spacing = None
        try:
            if style.paragraph_format and style.paragraph_format.line_spacing:
                from docx.shared import Length
                sp = style.paragraph_format.line_spacing
                if isinstance(sp, (int, float)) and sp > 3:
                    style_line_spacing = round(float(sp), 2)
                elif isinstance(sp, (int, float)):
                    style_line_spacing = round(float(sp) * 16, 2)
                else:
                    style_line_spacing = round(Length(sp, 0).pt, 2)
        except Exception:
            pass

        # 读取样式的对齐
        style_alignment = None
        try:
            if style.paragraph_format and style.paragraph_format.alignment is not None:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                _map = {
                    WD_ALIGN_PARAGRAPH.LEFT: "left",
                    WD_ALIGN_PARAGRAPH.CENTER: "center",
                    WD_ALIGN_PARAGRAPH.RIGHT: "right",
                    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
                }
                style_alignment = _map.get(style.paragraph_format.alignment)
        except Exception:
            pass

        # 读取样式的首行缩进
        style_indent = None
        try:
            if style.paragraph_format and style.paragraph_format.first_line_indent:
                from docx.shared import Length as L
                style_indent = round(L(style.paragraph_format.first_line_indent, 0).pt, 1)
        except Exception:
            pass

        # 应用 fallback：run 没有直接格式时用样式值
        for run in runs:
            if run.format.font_size_pt is None and style_font_size:
                run.format.font_size_pt = style_font_size
            if not run.format.font_name and style_font_name:
                run.format.font_name = style_font_name

        # 段落级格式 fallback
        if para_format.line_spacing_pt is None and style_line_spacing:
            para_format.line_spacing_pt = style_line_spacing
        if para_format.alignment is None and style_alignment:
            para_format.alignment = style_alignment
        if para_format.first_line_indent_pt is None and style_indent:
            para_format.first_line_indent_pt = style_indent

    except Exception:
        pass


def _parse_paragraph_format(para) -> ParagraphFormat:
    """Parse paragraph-level formatting comprehensively."""
    pf = para.paragraph_format

    # 对齐
    alignment = "left"
    if para.alignment is not None:
        alignment_map = {
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
        }
        alignment = alignment_map.get(para.alignment, "left")

    # 行间距
    line_spacing_pt = None
    line_spacing_rule = None
    if pf.line_spacing is not None:
        try:
            # 优先使用 line_spacing_rule 来区分模式
            from docx.enum.text import WD_LINE_SPACING
            rule = pf.line_spacing_rule
            if rule == WD_LINE_SPACING.MULTIPLE:
                # 倍数模式：pf.line_spacing 是倍数值 (如 1.5)
                # 转换为 pt：倍数 × 16pt（公文正文标准字号）
                line_spacing_pt = float(pf.line_spacing) * 16
                line_spacing_rule = "multiple"
            elif rule == WD_LINE_SPACING.AT_LEAST:
                # 最小值模式：pf.line_spacing 是 Length 对象
                line_spacing_pt = round(Length(pf.line_spacing, 0).pt, 2)
                line_spacing_rule = "atLeast"
            elif isinstance(pf.line_spacing, (int, float)):
                if pf.line_spacing > 3:
                    # 值大于3，视为固定行距（pt值）
                    line_spacing_pt = float(pf.line_spacing)
                    line_spacing_rule = "exact"
                else:
                    # 值<=3，视为倍数
                    line_spacing_pt = float(pf.line_spacing) * 16
                    line_spacing_rule = "multiple"
            else:
                # Length 对象（EXACTLY 模式或其他）
                line_spacing_pt = round(Length(pf.line_spacing, 0).pt, 2)
                line_spacing_rule = "exact"
        except Exception:
            pass

    # 首行缩进
    first_line_indent_pt = _safe_pt2(pf.first_line_indent)

    # XML 级别回退：python-docx 仅读取 w:firstLine，中文公文常用 w:firstLineChars
    # 当 python-docx 无法读取时（仅设了 firstLineChars），从 XML 直接解析
    if first_line_indent_pt is None:
        try:
            pPr = para._element.pPr
            if pPr is not None:
                ind = pPr.find(qn('w:ind'))
                if ind is not None:
                    chars_val = ind.get(qn('w:firstLineChars'))
                    if chars_val:
                        # firstLineChars 单位为百分之一字符，200 = 2字符
                        # 1字符 = 1字号 = 16pt（公文正文标准字号）
                        first_line_indent_pt = round(float(chars_val) / 100 * 16, 2)
        except Exception:
            pass

    left_indent_pt = _safe_pt2(pf.left_indent)
    right_indent_pt = _safe_pt2(pf.right_indent)

    return ParagraphFormat(
        alignment=alignment,
        line_spacing_pt=line_spacing_pt,
        line_spacing_rule=line_spacing_rule,
        first_line_indent_pt=first_line_indent_pt,
        left_indent_pt=left_indent_pt,
        right_indent_pt=right_indent_pt,
        space_before_pt=_safe_pt2(pf.space_before, 0),
        space_after_pt=_safe_pt2(pf.space_after, 0),
    )


# ---------------------------------------------------------------------------
#  Run
# ---------------------------------------------------------------------------

def _parse_run(run, index: int) -> Run:
    """Parse a single text run with full font information."""
    font = run.font

    font_size_pt = None
    if font.size:
        try:
            # font.size 返回 Length 对象（EMU 单位），用 .pt 属性转为磅值
            # 错误写法：Pt(font.size).pt → 会把 EMU 值当成 pt 返回（如 203200.0 而非 16.0）
            font_size_pt = round(font.size.pt, 1)
        except Exception:
            pass

    # 使用 font_utils 获取有效中文字体
    effective_font = get_effective_font(run)

    # 获取 RGB 颜色
    color_rgb = None
    if font.color and font.color.rgb:
        color_rgb = str(font.color.rgb)

    return Run(
        text=run.text,
        index=index,
        format=RunFormat(
            font_name=effective_font,
            font_size_pt=font_size_pt,
            bold=font.bold if font.bold is not None else False,
            italic=font.italic if font.italic is not None else False,
            underline=font.underline if font.underline is not None else False,
            color=color_rgb,
        ),
    )


# ---------------------------------------------------------------------------
#  Table
# ---------------------------------------------------------------------------

def _parse_table(table, index: int) -> Table:
    """Parse a table with cell-level paragraph preservation."""
    cells = []
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            # Parse cell paragraphs
            cell_paras = []
            for p_idx, para in enumerate(cell.paragraphs):
                cell_paras.append(_parse_paragraph(para, p_idx))

            cells.append(TableCell(
                text=cell.text,
                row=row_idx,
                col=col_idx,
                paragraphs=cell_paras,
            ))

    return Table(
        index=index,
        rows=len(table.rows),
        cols=len(table.columns) if table.rows else 0,
        cells=cells,
    )


# ---------------------------------------------------------------------------
#  Helpers
# ---------------------------------------------------------------------------

def _safe_pt2(value, default: float | None = None) -> float | None:
    """Safely convert a docx Length (EMU) to points float."""
    try:
        if value is None:
            return default
        # python-docx 属性返回 Length 对象（EMU），用 .pt 转为磅值
        # 注意：Pt(value).pt 不会做单位转换，会把 EMU 原值当 pt 返回（BUG）
        return round(Length(value, 0).pt, 2)
    except Exception:
        return default