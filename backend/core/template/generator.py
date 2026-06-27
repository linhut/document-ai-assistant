# This file is part of the Official Document AI Assistant.
# (c) 2026 Jose AI (https://www.linhut.cn)
# Licensed under the MIT License. See the LICENSE file for details.
"""
Style Generator: 将 YAML 模板转换为 Word 模板文件（.dotx / .docx）。

产物：
  .docx  — 预填充内容的模板文档（可直接使用）
  .dotx  — Word 模板文件（安装到模板库后可反复使用）

关键：
  - 所有字体必须通过 font_utils.set_run_font() 设置
  - 样式写入 Word 样式库
  - 用户新建文档时自动继承样式
"""
from __future__ import annotations
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from core.document.font_utils import set_run_font, LATIN_FONT
from core.template.style_manager import get_template
from utils.logger import logger


# 对齐方式映射
_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def generate_docx_template(template_id: str, output_path: Path | str,
                            content: dict | None = None) -> Path:
    """
    从 YAML 模板生成 .docx 模板文档。

    Args:
        template_id: 模板 ID（如 "notice"）
        output_path: 输出文件路径
        content: 可选的内容 dict，含 title/paragraphs 等

    Returns:
        输出文件路径
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    template = get_template(template_id)
    if not template:
        raise ValueError(f"Template not found: {template_id}")

    doc = Document()

    # 1. 应用页面设置
    _apply_page_setup(doc, template.get("page", {}))

    # 2. 创建 Word 样式
    _create_styles(doc, template.get("styles", {}))

    # 3. 添加示例内容（如果有）
    if content:
        _add_content(doc, content, template.get("styles", {}))
    else:
        _add_sample_content(doc, template)

    # 4. 保存
    doc.save(str(output_path))
    logger.info(f"Generated template document: {output_path}")
    return output_path


def generate_dotx_template(template_id: str, output_path: Path | str) -> Path:
    """
    从 YAML 模板生成 .dotx Word 模板文件。

    .dotx 是 Word 模板格式，用户安装后：
    - 新建文档时可选择此模板
    - 自动继承样式定义

    Args:
        template_id: 模板 ID
        output_path: 输出文件路径

    Returns:
        输出文件路径
    """
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    template = get_template(template_id)
    if not template:
        raise ValueError(f"Template not found: {template_id}")

    doc = Document()

    # 1. 应用页面设置
    _apply_page_setup(doc, template.get("page", {}))

    # 2. 创建 Word 样式（核心）
    _create_styles(doc, template.get("styles", {}))

    # 3. 添加示例内容（与 docx 模板一致，确保内容充实）
    _add_sample_content(doc, template)

    # 4. 保存为 dotx（template=True 标记）
    # python-docx 不直接支持 .dotx，先保存为 .docx 然后重命名
    temp_docx = output_path.with_suffix(".docx")
    doc.save(str(temp_docx))

    # 重命名为 .dotx 并修正 Content Type
    if output_path.suffix == ".dotx":
        import shutil
        shutil.move(str(temp_docx), str(output_path))
        _fix_dotx_content_type(output_path)
    else:
        output_path = temp_docx

    logger.info(f"Generated dotx template: {output_path}")
    return output_path


def _fix_dotx_content_type(dotx_path: Path) -> None:
    """
    修正 .dotx 文件的 Content Type。
    python-docx 生成的文件 Content_Types.xml 标记为 document，
    需要改为 template 才是合法的 .dotx 格式。
    """
    import zipfile
    import tempfile

    content_types_xml = '[Content_Types].xml'
    # document → template 的替换对
    replacements = [
        (
            'wordprocessingml.document.main+xml',
            'wordprocessingml.template.main+xml',
        ),
    ]

    tmp_path = dotx_path.with_suffix('.dotx.tmp')
    try:
        with zipfile.ZipFile(dotx_path, 'r') as zin:
            with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)
                    if item.filename == content_types_xml:
                        text = data.decode('utf-8')
                        for old, new in replacements:
                            text = text.replace(old, new)
                        data = text.encode('utf-8')
                    zout.writestr(item, data)
        # 替换原文件
        import shutil
        shutil.move(str(tmp_path), str(dotx_path))
    except Exception as e:
        logger.warning(f"Failed to fix dotx content type: {e}")
        # 清理临时文件
        if tmp_path.exists():
            tmp_path.unlink()


# ---------------------------------------------------------------------------
#  Internal helpers
# ---------------------------------------------------------------------------

def _apply_page_setup(doc: Document, page_config: dict):
    """应用页面设置。"""
    if not page_config:
        return

    section = doc.sections[0]

    # 纸张大小
    size = page_config.get("size", "A4")
    if size == "A4":
        section.page_width = Mm(210)
        section.page_height = Mm(297)
    elif "width_mm" in page_config:
        section.page_width = Mm(page_config["width_mm"])
        section.page_height = Mm(page_config.get("height_mm", 297))

    # 页边距
    margins = page_config.get("margins", {})
    if "top" in margins:
        section.top_margin = _parse_margin(margins["top"])
    if "bottom" in margins:
        section.bottom_margin = _parse_margin(margins["bottom"])
    if "left" in margins:
        section.left_margin = _parse_margin(margins["left"])
    if "right" in margins:
        section.right_margin = _parse_margin(margins["right"])


def _create_styles(doc: Document, styles_config: dict):
    """
    在 Word 文档中创建公文样式。

    关键：使用 Word 样式库而非内联格式，
    这样用户新建文档时可直接使用这些样式。
    """
    for style_key, style_def in styles_config.items():
        style_name = style_def.get("style_name", f"公文{style_key}")

        # 检查样式是否已存在
        existing = None
        for s in doc.styles:
            if s.name == style_name:
                existing = s
                break

        if existing:
            _apply_style(existing, style_def)
        else:
            # 创建新样式
            try:
                new_style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
                _apply_style(new_style, style_def)
            except Exception as e:
                logger.warning(f"Failed to create style '{style_name}': {e}")


def _apply_style(style, style_def: dict):
    """将样式定义应用到 Word 样式对象。"""
    font_ea = style_def.get("font_east_asia", "")
    font_latin = style_def.get("font_latin", LATIN_FONT)
    size_str = style_def.get("size", "16pt")
    bold = style_def.get("bold", False)
    align_str = style_def.get("alignment", "justify")

    # 解析字号
    size_pt = float(str(size_str).replace("pt", "").strip())

    # 设置字体（通过 OXML 写入 eastAsia）
    font = style.font
    font.name = font_latin
    font.size = Pt(size_pt)
    font.bold = bold

    # 写入 eastAsia
    if font_ea:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        rPr = style.element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)

        rFonts.set(qn('w:ascii'), font_latin)
        rFonts.set(qn('w:hAnsi'), font_latin)
        rFonts.set(qn('w:eastAsia'), font_ea)
        rFonts.set(qn('w:cs'), font_ea)

    # 对齐
    pf = style.paragraph_format
    align = _ALIGN_MAP.get(align_str)
    if align is not None:
        pf.alignment = align

    # 行距
    ls_str = style_def.get("line_spacing")
    if ls_str:
        ls_pt = float(str(ls_str).replace("pt", "").strip())
        pf.line_spacing = Pt(ls_pt)

    # 首行缩进
    indent_str = style_def.get("first_line_indent")
    if indent_str:
        if "em" in str(indent_str):
            em_val = float(str(indent_str).replace("em", "").strip())
            pf.first_line_indent = Pt(em_val * 16)
        else:
            pf.first_line_indent = Pt(float(str(indent_str).replace("pt", "").strip()))

    # 段后间距
    after_str = style_def.get("space_after")
    if after_str:
        pf.space_after = Pt(float(str(after_str).replace("pt", "").strip()))


def _add_content(doc: Document, content: dict, styles_config: dict):
    """添加用户提供的内容到文档。"""
    # 标题
    title_text = content.get("title", "")
    if title_text:
        title_style = styles_config.get("title", {}).get("style_name", "公文标题")
        para = doc.add_paragraph()
        run = para.add_run(title_text)
        # 查找并应用样式
        for s in doc.styles:
            if s.name == title_style:
                para.style = s
                break
        set_run_font(run, styles_config.get("title", {}).get("font_east_asia", "方正小标宋简体"))

    # 段落
    body_style = styles_config.get("body", {}).get("style_name", "公文正文")
    for para_text in content.get("paragraphs", []):
        para = doc.add_paragraph()
        run = para.add_run(para_text)
        for s in doc.styles:
            if s.name == body_style:
                para.style = s
                break
        set_run_font(run, styles_config.get("body", {}).get("font_east_asia", "仿宋_GB2312"))


def _add_sample_content(doc: Document, template: dict):
    """添加示例内容。"""
    name = template.get("name", "文档")
    styles = template.get("styles", {})
    body_font = styles.get("body", {}).get("font_east_asia", "仿宋_GB2312")
    title_font = styles.get("title", {}).get("font_east_asia", "方正小标宋简体")

    # 标题
    title_style_name = styles.get("title", {}).get("style_name", "公文标题")
    para = doc.add_paragraph()
    run = para.add_run(f"关于XXX的{name}")
    for s in doc.styles:
        if s.name == title_style_name:
            para.style = s
            break
    set_run_font(run, title_font)

    # 正文
    body_style_name = styles.get("body", {}).get("style_name", "公文正文")
    samples = [
        "各部门、各单位：",
        "",
        "根据XXX要求，现将有关事项通知如下：",
        "一、工作目标",
        "（具体内容）",
        "二、工作要求",
        "（具体内容）",
        "",
        "（单位名称）",
        "XXXX年XX月XX日",
    ]

    for text in samples:
        para = doc.add_paragraph()
        if text:
            run = para.add_run(text)
            for s in doc.styles:
                if s.name == body_style_name:
                    para.style = s
                    break
            set_run_font(run, body_font)


def _parse_margin(value) -> 'Mm':
    """解析页边距值。"""
    value = str(value).strip()
    if "cm" in value:
        return Mm(float(value.replace("cm", "").strip()) * 10)
    if "mm" in value:
        return Mm(float(value.replace("mm", "").strip()))
    return Mm(float(value))