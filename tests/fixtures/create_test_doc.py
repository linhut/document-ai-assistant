"""
Create a simple test document for Document Engine testing.
"""
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def create_test_notice():
    """Create a simple notice document for testing."""
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Mm(210)  # A4
    section.page_height = Mm(297)
    section.top_margin = Mm(37)
    section.bottom_margin = Mm(35)
    section.left_margin = Mm(28)
    section.right_margin = Mm(26)

    # Title
    title = doc.add_heading('关于XXX工作的通知', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = '方正小标宋简体'
        run.font.size = Pt(22)

    # Recipient
    recipient = doc.add_paragraph('各部门、各单位：')
    recipient.paragraph_format.first_line_indent = Pt(0)
    for run in recipient.runs:
        run.font.name = '仿宋_GB2312'
        run.font.size = Pt(16)

    # Body
    body1 = doc.add_paragraph('根据上级有关文件精神，现就XXX工作通知如下：')
    body1.paragraph_format.first_line_indent = Pt(32)  # 2字符
    for run in body1.runs:
        run.font.name = '仿宋_GB2312'
        run.font.size = Pt(16)

    body2 = doc.add_paragraph('一、工作内容')
    for run in body2.runs:
        run.font.name = '仿宋_GB2312'
        run.font.size = Pt(16)
        run.font.bold = True

    body3 = doc.add_paragraph('请各单位认真做好相关工作，确保任务落实到位。')
    body3.paragraph_format.first_line_indent = Pt(32)
    for run in body3.runs:
        run.font.name = '仿宋_GB2312'
        run.font.size = Pt(16)

    # Signature
    signature = doc.add_paragraph('XXX单位')
    signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in signature.runs:
        run.font.name = '仿宋_GB2312'
        run.font.size = Pt(16)

    # Date
    date = doc.add_paragraph('2026年6月23日')
    date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in date.runs:
        run.font.name = '仿宋_GB2312'
        run.font.size = Pt(16)

    # Add a simple table
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'

    # Fill table
    headers = ['项目', '负责人', '完成时间']
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.name = '仿宋_GB2312'
                run.font.size = Pt(14)

    data = [
        ['任务1', '张三', '7月1日'],
        ['任务2', '李四', '7月15日'],
    ]
    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, value in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = value

    # Header
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = 'XXX单位文件'
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in header_para.runs:
        run.font.name = '仿宋_GB2312'
        run.font.size = Pt(14)

    # Footer
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = '第 1 页 共 1 页'
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_para.runs:
        run.font.name = '仿宋_GB2312'
        run.font.size = Pt(12)

    return doc


if __name__ == '__main__':
    # Create fixtures directory
    fixtures_dir = Path(__file__).parent
    fixtures_dir.mkdir(parents=True, exist_ok=True)

    # Create test document
    doc = create_test_notice()
    output_path = fixtures_dir / 'test_notice.docx'
    doc.save(str(output_path))

    print(f'Created test document: {output_path}')
    print(f'File size: {output_path.stat().st_size} bytes')
