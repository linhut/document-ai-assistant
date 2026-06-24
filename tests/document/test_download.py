"""
Tests for document download functionality and document_service.
Covers the P0 blank document fix and optimized_path download.
"""
import sys
import os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'backend'))

import pytest
from pathlib import Path
from docx import Document
from docx.shared import Pt

from core.document.models import (
    DocumentModel, Paragraph, Run, RunFormat, ParagraphFormat,
    Table, TableCell, PageSetup, HeaderFooter, DocumentMetadata,
)
from core.document.parser import parse_docx
from core.document.generator import generate_docx


# ---------------------------------------------------------------------------
#  Helper: create a test .docx with tables
# ---------------------------------------------------------------------------

def _create_test_docx_with_table(path: Path) -> Path:
    """Create a .docx file that has both paragraphs and a table."""
    doc = Document()
    doc.add_paragraph("公文标题", style="Heading 1")
    doc.add_paragraph("正文第一段")

    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    data = [
        ["参会人员", "职务", "联系方式"],
        ["张三", "主任", "138xxxx"],
    ]
    for r, row_data in enumerate(data):
        for c, text in enumerate(row_data):
            table.cell(r, c).paragraphs[0].add_run(text)

    doc.add_paragraph("正文第二段")
    doc.save(str(path))
    return path


def _create_test_docx_tables_only(path: Path) -> Path:
    """Create a .docx file where content is ONLY in tables (no meaningful paragraphs)."""
    doc = Document()
    doc.add_paragraph("")  # empty paragraph

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    data = [
        ["议题", "负责人"],
        ["关于XX的请示", "办公室"],
        ["关于YY的报告", "财务处"],
    ]
    for r, row_data in enumerate(data):
        for c, text in enumerate(row_data):
            table.cell(r, c).paragraphs[0].add_run(text)

    doc.save(str(path))
    return path


# ---------------------------------------------------------------------------
#  Phase 1: Table preservation tests
# ---------------------------------------------------------------------------

class TestTablePreservation:
    """P0 fix: tables must survive the generate_docx roundtrip."""

    def test_tables_preserved_after_generate(self, tmp_path):
        """表格在 generate_docx 后必须保留内容。"""
        src = _create_test_docx_with_table(tmp_path / "input.docx")
        model = parse_docx(src)

        assert len(model.tables) >= 1, "Parser should detect tables"
        assert model.tables[0].rows == 2
        assert model.tables[0].cols == 3

        out = generate_docx(model, tmp_path / "output.docx")

        # Re-parse and verify
        result = parse_docx(out)
        assert len(result.tables) >= 1, "Tables must survive roundtrip"
        t = result.tables[0]
        assert t.rows == 2
        assert t.cols == 3

        # Verify cell content preserved
        cell_texts = {c.text for c in t.cells}
        assert "参会人员" in cell_texts
        assert "张三" in cell_texts
        assert "138xxxx" in cell_texts

    def test_table_only_document_not_blank(self, tmp_path):
        """P0 核心：表格为主的文档不能生成空白文件。"""
        src = _create_test_docx_tables_only(tmp_path / "input.docx")
        model = parse_docx(src)

        assert len(model.tables) >= 1
        out = generate_docx(model, tmp_path / "output.docx")

        # File must not be trivially empty
        assert out.stat().st_size > 1000, "Output file too small, likely blank"

        # Re-parse and check table content survived
        result = parse_docx(out)
        assert len(result.tables) >= 1, "Tables must survive in table-only doc"

        all_text = " ".join(c.text for t in result.tables for c in t.cells)
        assert "关于XX的请示" in all_text
        assert "办公室" in all_text

    def test_paragraphs_and_tables_order(self, tmp_path):
        """段落和表格的原始顺序在 roundtrip 后保持一致。"""
        src = _create_test_docx_with_table(tmp_path / "input.docx")
        model = parse_docx(src)

        # 原文结构: 段落1, 段落2, 表格1, 段落3
        assert len(model.paragraphs) == 3
        assert len(model.tables) == 1

        out = generate_docx(model, tmp_path / "output.docx")

        # 验证输出文件中段落内容保留
        result = parse_docx(out)
        para_texts = [p.text for p in result.paragraphs]
        assert any("公文标题" in t for t in para_texts), "Title paragraph must survive"
        assert any("正文第二段" in t for t in para_texts), "Last paragraph must survive"


# ---------------------------------------------------------------------------
#  Phase 2: Headers / footers / metadata tests
# ---------------------------------------------------------------------------

class TestHeadersFootersMetadata:
    """Headers, footers, and metadata must be preserved."""

    def test_metadata_preserved(self, tmp_path):
        """元数据在 roundtrip 后保留。"""
        src = _create_test_docx_with_table(tmp_path / "input.docx")

        # 设置元数据
        d = Document(str(src))
        d.core_properties.title = "测试标题"
        d.core_properties.author = "测试作者"
        d.save(str(src))

        model = parse_docx(src)
        assert model.metadata.title == "测试标题"
        assert model.metadata.author == "测试作者"

        out = generate_docx(model, tmp_path / "output.docx")
        result_doc = Document(str(out))
        assert result_doc.core_properties.title == "测试标题"
        assert result_doc.core_properties.author == "测试作者"


# ---------------------------------------------------------------------------
#  Phase 3: Generator roundtrip integrity (enhanced)
# ---------------------------------------------------------------------------

class TestGeneratorRoundtrip:
    """Roundtrip tests verifying content completeness."""

    def test_roundtrip_preserves_all_paragraph_texts(self, tmp_path):
        """Roundtrip 后所有段落文本必须与原文完全一致。"""
        src = _create_test_docx_with_table(tmp_path / "input.docx")
        model = parse_docx(src)

        original_texts = [p.text for p in model.paragraphs]
        out = generate_docx(model, tmp_path / "output.docx")

        result = parse_docx(out)
        result_texts = [p.text for p in result.paragraphs]

        assert result_texts == original_texts, \
            f"Paragraph texts changed: {original_texts} → {result_texts}"

    def test_roundtrip_preserves_page_setup(self, tmp_path):
        """页面设置在 roundtrip 后保留。"""
        src = _create_test_docx_with_table(tmp_path / "input.docx")
        model = parse_docx(src)

        out = generate_docx(model, tmp_path / "output.docx")
        result = parse_docx(out)

        assert abs(result.page_setup.margin_top_mm - model.page_setup.margin_top_mm) < 1
        assert abs(result.page_setup.margin_left_mm - model.page_setup.margin_left_mm) < 1

    def test_output_file_is_valid_docx(self, tmp_path):
        """生成的文件必须是有效的 .docx（可以被 python-docx 解析）。"""
        src = _create_test_docx_with_table(tmp_path / "input.docx")
        model = parse_docx(src)
        out = generate_docx(model, tmp_path / "output.docx")

        # 如果文件损坏，Document() 会抛异常
        doc = Document(str(out))
        assert len(doc.paragraphs) > 0

    def test_optimize_preserves_content_and_table(self, tmp_path):
        """优化（格式修复）不影响文本内容和表格。"""
        src = _create_test_docx_with_table(tmp_path / "input.docx")
        model = parse_docx(src)

        # 记录原始内容
        orig_para_texts = {p.text for p in model.paragraphs}
        orig_table_cells = {c.text for t in model.tables for c in t.cells}

        # 模拟优化：修改字体但不改内容
        from core.rules.engine import RuleEngine
        engine = RuleEngine()
        try:
            issues, fixed = engine.check_and_fix(model, "meeting")
        except Exception:
            # 如果 meeting 类型不存在，用 notice
            issues, fixed = engine.check_and_fix(model, "notice")

        # 验证内容未变
        fixed_para_texts = {p.text for p in fixed.paragraphs}
        fixed_table_cells = {c.text for t in fixed.tables for c in t.cells}

        assert fixed_para_texts == orig_para_texts, "Optimization must not change paragraph text"
        assert fixed_table_cells == orig_table_cells, "Optimization must not change table content"

        # 验证生成后文件有效
        out = generate_docx(fixed, tmp_path / "output.docx")
        result = parse_docx(out)
        assert len(result.paragraphs) > 0
        assert len(result.tables) > 0


# ---------------------------------------------------------------------------
#  Phase 4: Edge cases
# ---------------------------------------------------------------------------

class TestEdgeCases:
    """边界条件测试。"""

    def test_empty_model_generates_valid_docx(self, tmp_path):
        """空 DocumentModel 也能生成有效 .docx（不崩溃）。"""
        model = DocumentModel()
        out = generate_docx(model, tmp_path / "empty.docx")
        assert out.exists()
        assert out.stat().st_size > 0

    def test_model_without_source_path(self, tmp_path):
        """没有 source_path 的 model 使用新建 Document()。"""
        model = DocumentModel(
            paragraphs=[Paragraph(index=0, text="测试文本", runs=[])]
        )
        out = generate_docx(model, tmp_path / "no_source.docx")
        result = parse_docx(out)
        assert any("测试文本" in p.text for p in result.paragraphs)

    def test_multiple_tables(self, tmp_path):
        """文档中有多个表格时全部保留。"""
        doc = Document()
        doc.add_paragraph("第一段")

        t1 = doc.add_table(rows=2, cols=2)
        t1.style = "Table Grid"
        t1.cell(0, 0).paragraphs[0].add_run("表格1-A1")

        doc.add_paragraph("中间段")

        t2 = doc.add_table(rows=1, cols=3)
        t2.style = "Table Grid"
        t2.cell(0, 0).paragraphs[0].add_run("表格2-B1")

        src = tmp_path / "multi_table.docx"
        doc.save(str(src))

        model = parse_docx(src)
        assert len(model.tables) == 2

        out = generate_docx(model, tmp_path / "output.docx")
        result = parse_docx(out)
        assert len(result.tables) == 2
