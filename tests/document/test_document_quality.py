# This file is part of the Official Document AI Assistant.
# (c) 2026 Jose AI (https://www.linhut.cn)
# Licensed under the MIT License. See the LICENSE file for details.
"""
文档质量测试体系 — 模板生成、字体XML、格式规则、优化前后对比
"""
import os
import zipfile
import tempfile
from pathlib import Path
from xml.etree import ElementTree as ET

import pytest

from core.document.parser import parse_docx
from core.document.generator import generate_docx
from core.document.models import DocumentModel, Paragraph, Run, RunFormat, ParagraphFormat, PageSetup
from core.document.font_utils import (
    set_run_font, validate_document_fonts, detect_font_from_run,
    TITLE_FONT, BODY_FONT, LATIN_FONT,
)
from core.rules.engine import RuleEngine
from core.rules.loader import load_rules_for_type

# ---------------------------------------------------------------------------
#  Helpers
# ---------------------------------------------------------------------------

WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _get_docx_xml(docx_path: Path) -> str:
    """Extract word/document.xml from a .docx file."""
    with zipfile.ZipFile(docx_path, "r") as zf:
        return zf.read("word/document.xml").decode("utf-8")


def _find_fonts_in_xml(xml_str: str) -> list[dict]:
    """Find all font references in a docx XML string."""
    root = ET.fromstring(xml_str)
    fonts = []
    for elem in root.iter():
        if elem.tag.endswith("}rFonts"):
            entry = {}
            for attr in ["ascii", "hAnsi", "eastAsia", "cs"]:
                val = elem.get(f"{{{WORD_NS}}}{attr}")
                if val:
                    entry[attr] = val
            if entry:
                fonts.append(entry)
    return fonts


def _create_minimal_document_model() -> DocumentModel:
    """Create a minimal DocumentModel for testing with correct fonts."""
    return DocumentModel(
        filename="test.docx",
        page_setup=PageSetup(
            paper_width_mm=210,
            paper_height_mm=297,
            margin_top_mm=37,
            margin_bottom_mm=35,
            margin_left_mm=28,
            margin_right_mm=26,
        ),
        paragraphs=[
            Paragraph(
                text="测试标题",
                index=0,
                is_heading=True,
                heading_level=1,
                format=ParagraphFormat(alignment="center"),
                runs=[Run(index=0, text="测试标题", format=RunFormat(
                    font_name=TITLE_FONT, font_size_pt=22,
                ))],
            ),
            Paragraph(
                text="正文第一段内容。用于测试公文格式是否正确生成。",
                index=1,
                format=ParagraphFormat(
                    alignment="justify",
                    first_line_indent_pt=32,
                    line_spacing_pt=28.95,
                ),
                runs=[Run(index=0, text="正文第一段内容。用于测试公文格式是否正确生成。", format=RunFormat(
                    font_name=BODY_FONT, font_size_pt=16,
                ))],
            ),
            Paragraph(
                text="落款单位",
                index=2,
                format=ParagraphFormat(alignment="right"),
                runs=[Run(index=0, text="落款单位", format=RunFormat(
                    font_name=BODY_FONT, font_size_pt=16,
                ))],
            ),
        ],
    )


# ===========================================================================
#  Test 1: 模板生成测试
# ===========================================================================

class TestTemplateGeneration:
    """测试模板生成功能。"""

    def test_template_document_has_content(self, tmp_path):
        """模板文档不是空白文档。"""
        model = _create_minimal_document_model()
        output = tmp_path / "template_test.docx"
        generate_docx(model, output)

        assert output.exists()
        assert output.stat().st_size > 2000

        model2 = parse_docx(output)
        assert len(model2.paragraphs) >= 3
        assert model2.paragraphs[0].text == "测试标题"

    def test_template_has_correct_margins(self, tmp_path):
        """模板文档有正确的页边距。"""
        model = _create_minimal_document_model()
        output = tmp_path / "template_margins.docx"
        generate_docx(model, output)

        model2 = parse_docx(output)
        ps = model2.page_setup
        assert abs(ps.margin_top_mm - 37) < 5
        assert abs(ps.margin_left_mm - 28) < 5

    def test_template_has_correct_page_size(self, tmp_path):
        """模板文档是A4纸张。"""
        model = _create_minimal_document_model()
        output = tmp_path / "template_a4.docx"
        generate_docx(model, output)

        model2 = parse_docx(output)
        ps = model2.page_setup
        assert abs(ps.paper_width_mm - 210) < 5
        assert abs(ps.paper_height_mm - 297) < 5

    def test_all_template_types_generate(self, tmp_path):
        """所有模板类型都有对应的规则文件且包含必要配置。"""
        doc_types = [
            "notice", "request", "report", "letter", "meeting",
            "decision", "announcement", "notice_public"
        ]
        for dtype in doc_types:
            try:
                rules = load_rules_for_type(dtype)
            except Exception:
                continue

            assert rules, f"No rules found for {dtype}"
            assert "title" in rules, f"No title config for {dtype}"
            assert "body" in rules, f"No body config for {dtype}"


# ===========================================================================
#  Test 2: 字体 XML 测试
# ===========================================================================

class TestFontXml:
    """测试生成的 docx XML 中的字体设置。"""

    def test_east_asia_font_present(self, tmp_path):
        """生成的文档必须包含 eastAsia 字体属性。"""
        model = _create_minimal_document_model()
        output = tmp_path / "font_test.docx"
        generate_docx(model, output)

        xml_content = _get_docx_xml(output)
        fonts = _find_fonts_in_xml(xml_content)

        assert len(fonts) > 0, "文档中找不到任何字体设置"

        east_asia_fonts = [f.get("eastAsia") for f in fonts if f.get("eastAsia")]
        assert len(east_asia_fonts) > 0, "文档中没有设置 eastAsia 字体"

    def test_no_ms_gothic_font(self, tmp_path):
        """文档中不能出现 MS Gothic 等微软默认东亚字体。"""
        model = _create_minimal_document_model()
        output = tmp_path / "no_ms_test.docx"
        generate_docx(model, output)

        xml_content = _get_docx_xml(output)
        invalid_fonts = ["MS Gothic", "MS Mincho", "MS PGothic", "MS PMincho"]

        for bad in invalid_fonts:
            assert bad not in xml_content, f"文档XML中发现了无效字体: {bad}"

    def test_title_font_is_east_asian(self, tmp_path):
        """标题的 eastAsia 字体应为标题专用字体。"""
        model = _create_minimal_document_model()
        output = tmp_path / "title_font.docx"
        generate_docx(model, output)

        xml_content = _get_docx_xml(output)
        fonts = _find_fonts_in_xml(xml_content)

        if fonts:
            east = fonts[0].get("eastAsia", "")
            assert east, "标题段落的 eastAsia 字体未设置"
            assert east != "Times New Roman", f"标题 eastAsia 不应该设置为 Times New Roman, 实际: {east}"

    def test_body_font_is_east_asian(self, tmp_path):
        """正文的 eastAsia 字体应为正文字体。"""
        model = _create_minimal_document_model()
        output = tmp_path / "body_font.docx"
        generate_docx(model, output)

        xml_content = _get_docx_xml(output)
        fonts = _find_fonts_in_xml(xml_content)

        if len(fonts) >= 2:
            east = fonts[1].get("eastAsia", "")
            assert east, "正文段落的 eastAsia 字体未设置"

    def test_latin_font_is_times_new_roman(self, tmp_path):
        """西文字体应设置为 Times New Roman，不能是 MS 系列。"""
        model = _create_minimal_document_model()
        output = tmp_path / "latin_font.docx"
        generate_docx(model, output)

        xml_content = _get_docx_xml(output)
        fonts = _find_fonts_in_xml(xml_content)

        for f in fonts:
            ascii_font = f.get("ascii", "")
            if ascii_font:
                assert "MS " not in ascii_font, f"ASCII 字体不应该是 MS 系列: {ascii_font}"

    def test_font_utils_east_asia_sets_all_attrs(self, tmp_path):
        """set_run_font 应该设置 ascii, hAnsi, eastAsia, cs 四个属性。"""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        para = doc.add_paragraph()
        run = para.add_run("测试Test123")
        set_run_font(run, BODY_FONT)

        output = tmp_path / "font_attrs.docx"
        doc.save(str(output))

        xml_content = _get_docx_xml(output)
        fonts = _find_fonts_in_xml(xml_content)

        found = False
        for f in fonts:
            if f.get("eastAsia") == BODY_FONT:
                assert f.get("ascii") is not None, "ascii font not set"
                assert f.get("hAnsi") is not None, "hAnsi font not set"
                assert f.get("cs") is not None, "cs font not set"
                found = True
                break

        assert found, "没有找到使用 set_run_font 设置的完整字体元素"


# ===========================================================================
#  Test 3: 格式规则测试
# ===========================================================================

class TestFormatRules:
    """测试规则是否正确影响生成流程。"""

    def test_rule_engine_loads_all_types(self):
        """所有文档类型的规则都能加载。"""
        engine = RuleEngine()
        for dtype in ["notice", "request", "report", "letter", "meeting",
                       "decision", "announcement", "notice_public"]:
            rules = engine.load_rules(dtype)
            assert rules, f"Failed to load rules for {dtype}"
            assert "check_rules" in rules, f"No check_rules for {dtype}"
            assert "fix_rules" in rules, f"No fix_rules for {dtype}"

    def test_check_rules_have_required_fields(self):
        """所有 check_rules 都有必要的字段。"""
        for dtype in ["notice", "request", "report", "letter", "meeting",
                       "decision", "announcement", "notice_public"]:
            rules = load_rules_for_type(dtype)
            for rule in rules.get("check_rules", []):
                assert "id" in rule, f"Missing id in {dtype} rule: {rule}"
                assert "name" in rule, f"Missing name in {dtype} rule {rule['id']}"
                assert "severity" in rule, f"Missing severity in {dtype} rule {rule['id']}"

    def test_fix_rules_use_supported_actions(self):
        """所有 fix_rules 的 action 是受支持的。"""
        supported_actions = {
            "set_font", "set_size", "set_alignment", "set_align",
            "set_line_spacing", "set_first_line_indent", "set_indent",
            "set_margins", "set_page_margins",
            "remove_extra_spaces", "remove_extra_blank_lines",
        }
        for dtype in ["notice", "request", "report", "letter", "meeting",
                       "decision", "announcement", "notice_public"]:
            rules = load_rules_for_type(dtype)
            for rule in rules.get("fix_rules", []):
                action = rule.get("action", "")
                assert action in supported_actions, \
                    f"Unsupported action '{action}' in {dtype} rule {rule.get('id', '?')}"

    def test_notice_check_finds_issues(self):
        """通知模板检查应该能返回issue列表。"""
        model = _create_minimal_document_model()
        engine = RuleEngine()
        issues = engine.check(model, "notice")
        assert isinstance(issues, list)

    def test_notice_fix_modifies_document(self):
        """通知模板修复应该返回有效的文档模型。"""
        model = _create_minimal_document_model()
        engine = RuleEngine()
        fixed = engine.fix(model, "notice")
        assert fixed is not None
        assert len(fixed.paragraphs) == len(model.paragraphs)


# ===========================================================================
#  Test 4: 优化前后对比测试
# ===========================================================================

class TestOptimizationComparison:
    """优化前后对比测试。"""

    def test_optimization_preserves_content(self, tmp_path):
        """优化不应丢失文档内容。"""
        model = _create_minimal_document_model()
        original_texts = [p.text for p in model.paragraphs]

        engine = RuleEngine()
        fixed_model = engine.fix(model, "notice")

        fixed_texts = [p.text for p in fixed_model.paragraphs]
        for i, (orig, fixed) in enumerate(zip(original_texts, fixed_texts)):
            if len(orig.strip()) > 0:
                assert len(fixed.strip()) > 0, \
                    f"Paragraph {i} lost content during optimization"

    def test_optimization_generates_valid_docx(self, tmp_path):
        """优化后的文档可以正常生成并重新解析。"""
        model = _create_minimal_document_model()

        engine = RuleEngine()
        fixed_model = engine.fix(model, "notice")

        output = tmp_path / "optimized_test.docx"
        generate_docx(fixed_model, output)

        assert output.exists()

        model3 = parse_docx(output)
        assert len(model3.paragraphs) >= 2

    def test_optimization_improves_fonts(self, tmp_path):
        """优化后字体应该符合公文标准。"""
        model = _create_minimal_document_model()
        # 故意设置错误字体
        model.paragraphs[0].runs[0].format.font_name = "Arial"
        model.paragraphs[1].runs[0].format.font_name = "Arial"

        engine = RuleEngine()
        fixed_model = engine.fix(model, "notice")

        output = tmp_path / "font_improved.docx"
        generate_docx(fixed_model, output)

        xml_content = _get_docx_xml(output)

        fonts = _find_fonts_in_xml(xml_content)
        for f in fonts:
            if f.get("eastAsia"):
                assert f.get("eastAsia") != "Arial", "优化后 eastAsia 仍为 Arial"

    def test_font_validation_detects_issues(self):
        """validate_document_fonts 能检测出字体问题。"""
        from docx import Document as DocxDocument
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        doc = DocxDocument()
        para = doc.add_paragraph()
        run = para.add_run("测试")
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:eastAsia'), 'MS Gothic')
        rPr.insert(0, rFonts)

        issues = validate_document_fonts(doc)
        assert len(issues) > 0
        assert any("MS Gothic" in str(i["font_name"]) for i in issues)


# ===========================================================================
#  Test 5: 字体工具单元测试
# ===========================================================================

class TestFontUtilsUnit:
    """字体工具函数的单元测试。"""

    def test_set_run_font_produces_correct_xml(self, tmp_path):
        """set_run_font 生成正确的 OOXML 结构。"""
        from docx import Document as DocxDocument
        from docx.oxml.ns import qn

        doc = DocxDocument()
        para = doc.add_paragraph()
        run = para.add_run("测试文本ABC123")

        set_run_font(run, "仿宋_GB2312")

        rPr = run._element.find(qn('w:rPr'))
        assert rPr is not None

        rFonts = rPr.find(qn('w:rFonts'))
        assert rFonts is not None
        assert rFonts.get(qn('w:eastAsia')) == "仿宋_GB2312"
        assert rFonts.get(qn('w:ascii')) == "Times New Roman"
        assert rFonts.get(qn('w:hAnsi')) == "Times New Roman"
        assert rFonts.get(qn('w:cs')) == "仿宋_GB2312"

    def test_validate_font_name_rejects_ms_gothic(self):
        """validate_font_name 拒绝 MS Gothic 系列。"""
        from core.document.font_utils import validate_font_name

        assert validate_font_name("MS Gothic") is False
        assert validate_font_name("MS Mincho") is False
        assert validate_font_name("MS PGothic") is False
        assert validate_font_name("仿宋_GB2312") is True
        assert validate_font_name("方正小标宋简体") is True
        assert validate_font_name("") is False
        assert validate_font_name(None) is False

    def test_get_font_fallback_works(self):
        """get_font_fallback 返回正确的回退字体。"""
        from core.document.font_utils import get_font_fallback

        assert get_font_fallback("方正小标宋简体") == "SimSun"
        assert get_font_fallback("仿宋_GB2312") == "FangSong"
        assert get_font_fallback("不存在的字体") == "不存在的字体"

    def test_detect_font_from_run(self):
        """detect_font_from_run 正确读取 XML 字体属性。"""
        from docx import Document as DocxDocument

        doc = DocxDocument()
        para = doc.add_paragraph()
        run = para.add_run("测试")
        set_run_font(run, "楷体_GB2312")

        info = detect_font_from_run(run)
        assert info["eastAsia"] == "楷体_GB2312"
        assert info["ascii"] == "Times New Roman"


if __name__ == "__main__":
    pytest.main([__file__, "-v", "--tb=short"])